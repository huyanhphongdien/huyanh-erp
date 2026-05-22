"""
Generate file .docx MẪU để test Auto-fill upload flow.

4 file output (giả lập 4 template Docs chưa gửi):
  - sample_SC_FOB.docx — Sales Contract FOB (chỉ {{contract_no}}, bank ghi "see CI")
  - sample_SC_CIF.docx — Sales Contract CIF (chỉ {{contract_no}}, bank ghi "see CI")
  - sample_PI_FOB.docx — Proforma Invoice FOB (đủ 6 token)
  - sample_PI_CIF.docx — Proforma Invoice CIF (đủ 6 token)

Chạy: python docs/contract-template-samples/generate_sample_templates.py
Output: docs/contract-template-samples/sample_*.docx

User test:
1. Upload 4 file này vào /sales/orders/new (Sale page) — upload flow
2. Đăng nhập Phú LV mở review
3. Gõ số HĐ + chọn Bank → bấm Auto-fill
4. Download file ② → kiểm 6 token đã thay đúng chưa
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
PRIMARY = RGBColor(0x1B, 0x4D, 0x3E)


def setup_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    return doc


def add_heading(doc, text, size=14, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = PRIMARY
    return p


def add_line(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_kv(doc, key, value, key_width=4):
    """Add 'Key: Value' line. Value có thể là token {{...}}."""
    p = doc.add_paragraph()
    run_k = p.add_run(f'{key.ljust(key_width * 4)} ')
    run_k.bold = True
    run_k.font.size = Pt(11)
    run_v = p.add_run(value)
    run_v.font.size = Pt(11)


# ═══════════════════════════════════════════════════════════════════════════════
# SC TEMPLATE (Sales Contract — chỉ contract_no, bank "see CI")
# ═══════════════════════════════════════════════════════════════════════════════

def build_sc(incoterm: str):
    doc = setup_doc()
    add_heading(doc, 'SALES CONTRACT')
    add_heading(doc, f'(SAMPLE — INCOTERM {incoterm} — for ERP Auto-fill testing)', size=10)

    p = doc.add_paragraph()
    p.add_run('No.: ').bold = True
    p.add_run('{{contract_no}}')

    add_line(doc, 'Date: 22 May 2026')
    add_line(doc, '')

    add_line(doc, 'BETWEEN:', bold=True)
    add_line(doc, 'Seller: HUY ANH RUBBER COMPANY LIMITED')
    add_line(doc, 'Address: Lot 17, Phong Dien Industrial Zone, Hue City, Vietnam')
    add_line(doc, '')
    add_line(doc, 'Buyer: APOLLO TYRES LTD (sample buyer)')
    add_line(doc, 'Address: ...')
    add_line(doc, '')

    add_line(doc, '1. COMMODITY:', bold=True)
    add_line(doc, 'Natural rubber SVR3L, Vietnam origin')
    add_line(doc, '')

    add_line(doc, '2. QUANTITY:', bold=True)
    add_line(doc, '42 MT (3 x 20\'), partial shipment allowed')
    add_line(doc, '')

    add_line(doc, '3. PRICE:', bold=True)
    add_line(doc, f'USD 2,350/MT {incoterm} Belawan Port, Indonesia (Incoterms 2020)')
    add_line(doc, '')

    add_line(doc, '4. PAYMENT:', bold=True)
    add_line(doc, 'Note: Payment to be made into the bank A/C stated in Commercial Invoice.', bold=False)
    add_line(doc, '')

    add_line(doc, '5. SHIPMENT TIME:', bold=True)
    add_line(doc, 'June 2026')
    add_line(doc, '')

    add_line(doc, '___________________________________________________________________')
    add_line(doc, 'SELLER\'s Signature                       BUYER\'s Signature')

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# PI TEMPLATE (Proforma Invoice — đủ 6 token)
# ═══════════════════════════════════════════════════════════════════════════════

def build_pi(incoterm: str):
    doc = setup_doc()
    add_heading(doc, 'PROFORMA INVOICE / COMMERCIAL INVOICE')
    add_heading(doc, f'(SAMPLE — INCOTERM {incoterm} — for ERP Auto-fill testing)', size=10)

    p = doc.add_paragraph()
    p.add_run('Contract No.: ').bold = True
    p.add_run('{{contract_no}}')

    add_line(doc, 'Date: 22 May 2026')
    add_line(doc, '')

    add_line(doc, 'SELLER:', bold=True)
    add_line(doc, 'HUY ANH RUBBER COMPANY LIMITED')
    add_line(doc, 'Lot 17, Phong Dien IZ, Hue City, Vietnam')
    add_line(doc, '')

    add_line(doc, 'BUYER:', bold=True)
    add_line(doc, 'APOLLO TYRES LTD (sample)')
    add_line(doc, '')

    add_line(doc, 'GOODS:', bold=True)
    add_line(doc, 'SVR3L Natural Rubber — 42 MT — USD 98,700.00')
    add_line(doc, f'Incoterm: {incoterm} Belawan Port, Indonesia')
    add_line(doc, '')

    add_line(doc, 'BANK DETAILS FOR PAYMENT:', bold=True)
    add_kv(doc, 'Account Name:', '{{bank_account_name}}')
    add_kv(doc, 'Account No.:', '{{bank_account_no}}')
    add_kv(doc, 'Bank:', '{{bank_full_name}}')
    add_kv(doc, 'Address:', '{{bank_address}}')
    add_kv(doc, 'SWIFT:', '{{bank_swift}}')
    add_line(doc, '')

    add_line(doc, '___________________________________________________________________')
    add_line(doc, 'For and on behalf of HUY ANH RUBBER COMPANY LIMITED')

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    files = [
        ('sample_SC_FOB.docx', build_sc('FOB')),
        ('sample_SC_CIF.docx', build_sc('CIF')),
        ('sample_PI_FOB.docx', build_pi('FOB')),
        ('sample_PI_CIF.docx', build_pi('CIF')),
    ]
    for name, doc in files:
        path = ROOT / name
        doc.save(str(path))
        print(f'OK Generated: {name} ({path.stat().st_size // 1024} KB)')

    print()
    print('=' * 60)
    print('CACH TEST:')
    print('=' * 60)
    print('1. Vao /sales/orders/new, upload 4 file nay (Sale page)')
    print('2. Dang nhap voi phulv@huyanhrubber.com')
    print('3. Vao /sales/contracts/review')
    print('4. Mo HD vua tao')
    print('5. Go So HD (vd: HA20260100) + chon Bank (vd: Vietin Hue)')
    print('6. Bam Auto-fill')
    print('7. Download tung file o o (R) -> kiem 6 token da thay dung chua')


if __name__ == '__main__':
    main()
