"""
Build docxtemplater templates from 4 sample contracts.

QUAN TRỌNG: bank info được placeholderize để Phú (Kiểm tra) nhập per-order
(do mỗi HĐ có thể dùng Vietcombank / Vietinbank / BIDV… khác nhau). 5 field:
{bank_account_name}, {bank_account_no}, {bank_full_name}, {bank_address}, {bank_swift}



Source: SC + PI ( CIF)/* và SC + PI (FOB)/*
Output: docs/contract-templates/template_{SC,PI}_{CIF,FOB}.docx

Cách hoạt động:
- Mở từng file mẫu bằng python-docx
- Với mỗi paragraph và mỗi cell trong table → áp dụng map replace
  (literal-string find → docxtemplater placeholder `{var}` hoặc `{#docs}…{/docs}`)
- Khi replace, gộp toàn bộ runs trong paragraph thành 1 run duy nhất
  (đảm bảo placeholder không bị Word split thành nhiều run, docxtemplater sẽ detect được)

Lưu ý:
- Format intra-paragraph (bold/italic giữa câu) có thể bị mất. OK với contract templates.
- Replacements dict được build từ preset trong mock (apollo/yoongdo) — value preset
  được thay bằng key placeholder.
"""

from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent.parent  # repo root
OUT_DIR = ROOT / "docs" / "contract-templates"
PUBLIC_DIR = ROOT / "public" / "contract-templates"
OUT_DIR.mkdir(parents=True, exist_ok=True)
PUBLIC_DIR.mkdir(parents=True, exist_ok=True)

# ----------------------------------------------------------------------
# Replacement maps cho từng file (literal → placeholder)
# Tham chiếu chính xác text trong file mẫu (lưu ý dấu cách lạ "2, 460", "49 ,593.60")
# ----------------------------------------------------------------------

# Bank info chung cho cả 4 file — Phú LV (Kiểm tra) nhập per-order
# 4 file mẫu đều dùng cùng 1 account Vietin Hue của Huy Anh; nhưng thực tế
# có thể đổi bank/branch/account → cần placeholder.
BANK_REPL = [
    ("ACCOUNT NAME: HUY ANH RUBBER COMPANY LIMITED", "ACCOUNT NAME: {bank_account_name}"),
    ("ACCOUNT NO: 111002648221", "ACCOUNT NO: {bank_account_no}"),
    ("AT BANK’S NAME: VIETNAM JOINT STOCK COMMERCIAL BANK FOR INDUSTRY AND TRADE HUE BRANCH",
     "AT BANK’S NAME: {bank_full_name}"),
    ("AT BANK'S NAME: VIETNAM JOINT STOCK COMMERCIAL BANK FOR INDUSTRY AND TRADE HUE BRANCH",
     "AT BANK’S NAME: {bank_full_name}"),
    ("ADDRESS: 02 LE QUY DON STREET, THUAN HOA WARD, HUE CITY, VIET NAM",
     "ADDRESS: {bank_address}"),
    ("SWIFT CODE: ICBVVNVX460", "SWIFT CODE: {bank_swift}"),
]

# CIF Yoong Do (SC)
REPL_SC_CIF = [
    ("HA20260053", "{contract_no}"),
    ("08 th May 2026", "{contract_date}"),
    ("08th May 2026", "{contract_date}"),
    ("YOONG DO ENGINEERING CO.,LTD", "{buyer_name}"),
    ("295, YANGYEON-RO, NAM-MYEON, YANGJU-SI, GYEONGGI-DO, REPUBLIC OF KOREA", "{buyer_address}"),
    ("NATURAL RUBBER SVR3L", "NATURAL RUBBER {grade}"),
    ("CIF – INCHEON , KOREA", "{incoterm} – {pod}"),
    ("CIF - INCHEON , KOREA", "{incoterm} – {pod}"),
    ("CIF – INCHEON, KOREA", "{incoterm} – {pod}"),
    ("Incheon , Korea", "{pod}"),
    ("Incheon, Korea", "{pod}"),
    ("Any port , Viet Nam", "{pol}"),
    ("Any port, Viet Nam", "{pol}"),
    # Numbers (đặt sau text dài hơn để tránh conflict)
    ("20.16", "{quantity}"),
    ("2, 460", "{unit_price}"),
    ("2,460", "{unit_price}"),
    ("49 ,593.60", "{amount}"),
    ("49,593.60", "{amount}"),
    # Packing — replace full hardcoded text bằng placeholder.
    # Source thực tế có ký tự NO-BREAK SPACE (\xa0) giữa "," và "35kg" + smart quote ’.
    ("In 1.26mts Per Wooden pallets, \xa035kg/bale with thick polybag", "{packing_desc}"),
    ("In 1.26mts Per Wooden pallets, 35kg/bale with thick polybag", "{packing_desc}"),
    ("In 1.26mts Per Wooden pallets , 35 kg/bale with thick polybag", "{packing_desc}"),
    ("1.26mts Per Wooden pallets , 35 kg/bale with thick polybag", "{packing_desc}"),
    # Container line — source có "01 x 20’" (smart quote = feet symbol).
    # Wrap "{pallets_total} Wooden pallets / " bằng {#has_pallets}...{/has_pallets}
    # để ẩn segment khi loose_bale (không có pallet).
    ("576 bales/ 16 Wooden pallets/ 01 x 20’", "{bales_total} bales / {#has_pallets}{pallets_total} Wooden pallets / {/has_pallets}{containers} x {cont_type}"),
    ("576 bales/ 16 Wooden pallets/ 01 x 20'", "{bales_total} bales / {#has_pallets}{pallets_total} Wooden pallets / {/has_pallets}{containers} x {cont_type}"),
    ("576 bales/ 16 Wooden pallets/ 01 x 20DC", "{bales_total} bales / {#has_pallets}{pallets_total} Wooden pallets / {/has_pallets}{containers} x {cont_type}"),
    ("576 bales/ 16 Wooden pallets/ 01 x 20", "{bales_total} bales / {#has_pallets}{pallets_total} Wooden pallets / {/has_pallets}{containers} x {cont_type}"),
    # Fumigation certificate — chỉ hiện khi packing=wooden_pallet (Korea/EU requirement).
    # Source thực: " Fumigation certificate" (leading space). CHỈ 1 pattern — pattern
    # thứ 2 "Fumigation certificate" sẽ fire LẦN NỮA trên text đã wrap → double-wrap.
    (" Fumigation certificate", "{#has_fumigation} Fumigation certificate{/has_fumigation}"),
    # L/C draft sentence — chỉ hiện khi payment là L/C (CIF Korea/EU mới có L/C).
    ("The L/C draft must be opened within five (5) days from the contract signing date.",
     "{#is_lc_payment}The L/C draft must be opened within five (5) days from the contract signing date.{/is_lc_payment}"),
    # Shipment
    ("Time of shipment: June , 2026", "Time of shipment: {shipment_time}"),
    ("Time of shipment: June, 2026", "Time of shipment: {shipment_time}"),
    ("Partial shipment: Not Allowed", "Partial shipment: {partial}"),
    ("Trans s hipment: Allowed", "Transshipment: {trans}"),
    ("Transshipment: Allowed", "Transshipment: {trans}"),
    # Payment
    ("LC at sight. The L/C draft must be opened within five (5) days from the contract signing date.",
     "{payment}. {payment_extra}"),
    ("LC at sight", "{payment}"),
    # Claims
    ("within 20 days of receipt of goods", "within {claims_days} days of receipt of goods"),
    # Freight
    ("freight prepaid", "{freight_mark}"),
    # Arbitration — Sale có thể chọn SICOM Singapore / LCIA London / ...
    ("by SICOM Singapore arbitration clause", "by {arbitration} arbitration clause"),
]
REPL_SC_CIF += BANK_REPL

# CIF Yoong Do (PI)
REPL_PI_CIF = [
    ("HA20260053", "{contract_no}"),
    ("08 th May 2026", "{contract_date}"),
    ("08th May 2026", "{contract_date}"),
    ("YOONG DO ENGINEERING CO.,LTD", "{buyer_name}"),
    ("295, YANGYEON-RO, NAM-MYEON, YANGJU-SI, GYEONGGI-DO, REPUBLIC OF KOREA", "{buyer_address}"),
    ("CIF – INCHEON , KOREA", "{incoterm} – {pod}"),
    ("CIF - INCHEON , KOREA", "{incoterm} – {pod}"),
    ("CIF – INCHEON, KOREA", "{incoterm} – {pod}"),
    ("NATURAL RUBBER SVR3L", "NATURAL RUBBER {grade}"),
    # Container — source thực: "01 x 20’ DC" (smart quote = feet, sau đó " DC")
    # Phải match cả block "20’ DC" để không thừa khi cont_type=20DC.
    ("01 x 20’ DC", "{containers} x {cont_type}"),
    ("01 x 20' DC", "{containers} x {cont_type}"),
    ("01 x 20DC’", "{containers} x {cont_type}"),
    ("01 x 20DC'", "{containers} x {cont_type}"),
    ("01 x 20DC", "{containers} x {cont_type}"),
    ("01 x 20", "{containers} x {cont_type}"),
    ("35kg/ bale with thick polybag – Wooden pallets", "{packing_desc}"),
    ("35kg/ bale with thick polybag - Wooden pallets", "{packing_desc}"),
    ("20.16", "{quantity}"),
    ("2, 460", "{unit_price}"),
    ("2,460", "{unit_price}"),
    ("49 ,593.60", "{amount}"),
    ("49,593.60", "{amount}"),
    ("Forty-Nine Thousand Five Hundred Ninety-Three US Dollars and Sixty Cents Only",
     "{amount_words}"),
    ("LC at sight. The L/C draft must be opened within five (5) days from the contract signing date.",
     "{payment}. {payment_extra}"),
    ("LC at sight", "{payment}"),
]
REPL_PI_CIF += BANK_REPL

# FOB Apollo (SC)
REPL_SC_FOB = [
    ("HA20260051", "{contract_no}"),
    ("08 th May 2026", "{contract_date}"),
    ("08th May 2026", "{contract_date}"),
    ("APOLLO TYRES LTD", "{buyer_name}"),
    ("7, INSTITUTIONAL AREA, SECTOR 32, GURGAON, INDIA 122001", "{buyer_address}"),
    ("1800 212 7070", "{buyer_phone}"),
    ("NATURAL RUBBER RSS3", "NATURAL RUBBER {grade}"),
    ("FOB DA NANG PORT, VIET NAM", "{incoterm} {pol}"),
    ("FOB DA NANG PORT , VIET NAM", "{incoterm} {pol}"),
    ("Da Nang port , Viet Nam", "{pol}"),
    ("Da Nang port, Viet Nam", "{pol}"),
    ("201 .6", "{quantity}"),
    ("201.6", "{quantity}"),
    ("2, 350", "{unit_price}"),
    ("2,350", "{unit_price}"),
    ("473 ,760.00", "{amount}"),
    ("473,760.00", "{amount}"),
    # Packing — source thực: "- 35kg/bales. Loose bales packing"
    ("35kg/bales. Loose bales packing", "{packing_desc}"),
    ("35kg /bales. Loose bales packing", "{packing_desc}"),
    ("35 kg/bales. Loose bales packing", "{packing_desc}"),
    # Container — source thực: "5,760 bales/ 10 x 20’" (smart quote)
    ("5,760 bales/ 10 x 20’", "{bales_total} bales / {containers} x {cont_type}"),
    ("5,760 bales/ 10 x 20'", "{bales_total} bales / {containers} x {cont_type}"),
    ("5,760 bales/ 10 x 20DC", "{bales_total} bales / {containers} x {cont_type}"),
    ("5,760 bales/ 10 x 20", "{bales_total} bales / {containers} x {cont_type}"),
    # Shipment (multi-lot) — nếu file gốc tách 2 paragraph "+ 1st Lot..." và "+ 2nd Lot...",
    # ta thay paragraph đầu thành "{shipment_time}" và xóa nội dung paragraph "+ 2nd Lot..."
    ("+ 1st Lot: Before 15th June, 2026 + 2nd Lot: Before 30th June, 2026", "{shipment_time}"),
    ("+ 1st Lot: Before 15th June, 2026", "{shipment_time}"),
    ("+ 2nd Lot: Before 30th June, 2026", ""),
    ("1st Lot: Before 15th June, 2026", "{shipment_time}"),
    # Single-line variants
    ("Partial shipment: Allowed", "Partial shipment: {partial}"),
    ("Trans s hipment: Allowed", "Transshipment: {trans}"),
    ("Transshipment: Allowed", "Transshipment: {trans}"),
    # Payment
    ("CAD 5 days", "{payment}"),
    ("within 20 days of receipt of goods", "within {claims_days} days of receipt of goods"),
    ("freight Collect", "{freight_mark}"),
    ("freight collect", "{freight_mark}"),
    # Arbitration — Sale có thể chọn SICOM Singapore / LCIA London / ...
    ("by SICOM Singapore arbitration clause", "by {arbitration} arbitration clause"),
]
REPL_SC_FOB += BANK_REPL

# FOB Apollo (PI)
REPL_PI_FOB = [
    ("HA20260051", "{contract_no}"),
    ("08 th May 2026", "{contract_date}"),
    ("08th May 2026", "{contract_date}"),
    ("APOLLO TYRES LTD", "{buyer_name}"),
    ("7, INSTITUTIONAL AREA, SECTOR 32, GURGAON, INDIA 122001", "{buyer_address}"),
    ("1800 212 7070", "{buyer_phone}"),
    ("FOB DA NANG PORT , VIET NAM", "{incoterm} {pol}"),
    ("FOB DA NANG PORT, VIET NAM", "{incoterm} {pol}"),
    ("NATURAL RUBBER RSS3", "NATURAL RUBBER {grade}"),
    # Container — source thực: "10 x 20’ DC" (smart quote + " DC")
    ("10 x 20’ DC", "{containers} x {cont_type}"),
    ("10 x 20' DC", "{containers} x {cont_type}"),
    ("10 x 20DC’", "{containers} x {cont_type}"),
    ("10 x 20DC'", "{containers} x {cont_type}"),
    ("10 x 20DC", "{containers} x {cont_type}"),
    ("10 x 20", "{containers} x {cont_type}"),
    ("35 kg/ bale . Loose bales packing", "{packing_desc}"),
    ("35 kg/ bale. Loose bales packing", "{packing_desc}"),
    ("201 .6", "{quantity}"),
    ("201.6", "{quantity}"),
    ("2, 350", "{unit_price}"),
    ("2,350", "{unit_price}"),
    ("473 ,760.00", "{amount}"),
    ("473,760.00", "{amount}"),
    ("Four Hundred Seventy-Three Thousand Seven Hundred Sixty US Dollars Only",
     "{amount_words}"),
    ("CAD 5 days", "{payment}"),
]
REPL_PI_FOB += BANK_REPL


def _replace_in_paragraph(paragraph, replacements):
    """
    Áp dụng list (find_str, replace_str) lên paragraph.
    Nếu có thay đổi: collapse toàn bộ runs thành 1 run duy nhất giữ format của run đầu.
    """
    original = paragraph.text
    if not original.strip():
        return False
    new_text = original
    for find, repl in replacements:
        if find in new_text:
            new_text = new_text.replace(find, repl)
    if new_text == original:
        return False

    runs = paragraph.runs
    if not runs:
        return False

    # Giữ run đầu tiên, đặt text mới, xóa các runs còn lại
    first_run = runs[0]
    first_run.text = new_text
    # Remove subsequent run elements
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    return True


def _iter_paragraphs(doc):
    """Yield all paragraphs in document (body + tables, recursive)."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                # Nested tables (rare but possible)
                for nested in cell.tables:
                    for nr in nested.rows:
                        for nc in nr.cells:
                            for np in nc.paragraphs:
                                yield np


def _normalize_head_letter(doc):
    """
    ★ NORMALIZE format head letter — KHÔNG xóa SELLER block, KHÔNG add letterhead. ★

    Page header (logo HUYANH + text letterhead nhỏ) đã có sẵn trong template gốc
    (rendered tự động ở đỉnh page). Em chỉ cần normalize format các paragraph
    body cho SC/PI sync nhau:

    - Title "SALES CONTRACT" / "PROFORMA INVOICE": EXACT match (không substring),
      sz=44 centered bold
    - "No:" / "Date:" lines: sz=28 centered bold + trim leading whitespace
    - SELLER block (THE SELLER + KHE MA + Tel/Fax): sz=26 (giữ alignment 'both')
    - KHÔNG ADD letterhead text (page header đã có)
    - KHÔNG XÓA SELLER block (mẫu có)
    """
    from docx.oxml.ns import qn

    def _set_align(p, val):
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = p._p.makeelement(qn('w:pPr'), {})
            p._p.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = pPr.makeelement(qn('w:jc'), {qn('w:val'): val})
            pPr.append(jc)
        else:
            jc.set(qn('w:val'), val)

    def _set_run_sz(run, sz_val):
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            rPr = run._r.makeelement(qn('w:rPr'), {})
            run._r.insert(0, rPr)
        for tag in (qn('w:sz'), qn('w:szCs')):
            el = rPr.find(tag)
            if el is None:
                el = rPr.makeelement(tag, {qn('w:val'): str(sz_val)})
                rPr.append(el)
            else:
                el.set(qn('w:val'), str(sz_val))

    def _set_bold(run, bold):
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            rPr = run._r.makeelement(qn('w:rPr'), {})
            run._r.insert(0, rPr)
        b = rPr.find(qn('w:b'))
        if bold and b is None:
            rPr.append(rPr.makeelement(qn('w:b'), {}))
        elif not bold and b is not None:
            rPr.remove(b)

    def _set_italic(run, italic):
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            rPr = run._r.makeelement(qn('w:rPr'), {})
            run._r.insert(0, rPr)
        it = rPr.find(qn('w:i'))
        if italic and it is None:
            rPr.append(rPr.makeelement(qn('w:i'), {}))
        elif not italic and it is not None:
            rPr.remove(it)

    TITLE_SET = {'SALES CONTRACT', 'SALE CONTRACT', 'PROFORMA INVOICE'}
    changed = 0
    for i, p in enumerate(doc.paragraphs[:12]):  # head letter trong 12 para đầu
        raw = p.text or ''
        text = raw.strip()
        upper = text.upper()
        # ─ Title — EXACT match (không substring!) — fix bug "This sales contract..."
        if upper in TITLE_SET:
            _set_align(p, 'center')
            for r in p.runs:
                _set_run_sz(r, 44)
                _set_bold(r, True)
            changed += 1
        # ─ No: HA... / Date: ... lines: RIGHT-aligned, italic + bold, sz=24 (12pt)
        # Match mẫu PI gốc của Apollo/Yoong Do (italic, right-aligned, không quá to)
        elif (text.startswith('No:') or text.startswith('No.:') or text.startswith('No ') or
              text.startswith('Date:') or text.startswith('Date ')):
            _set_align(p, 'right')
            for r in p.runs:
                _set_run_sz(r, 24)
                _set_bold(r, True)
                _set_italic(r, True)
            # Trim leading/trailing whitespace (PI_FOB src có "                     No: ..." padding)
            if p.runs:
                p.runs[0].text = p.runs[0].text.lstrip()
                p.runs[-1].text = p.runs[-1].text.rstrip()
            changed += 1
        # ─ SELLER block (THE SELLER + KHE MA + Tel/Fax): sz=26 (giữ alignment)
        elif ('THE SELLER' in upper or
              upper.startswith('KHE MA') or
              text.startswith('Tel:') or text.startswith('Tel ')):
            for r in p.runs:
                _set_run_sz(r, 26)
            changed += 1
    return changed


def _fix_pi_packing_in_cells(doc):
    """
    PI templates có dòng packing description trong cell sau "NATURAL RUBBER {grade}".
    PI_CIF source: "(35kg/ bale with thick polybag – Wooden pallets)" — 1 paragraph,
        replace patterns trong REPL_PI_CIF match → OK.
    PI_FOB source: "(35kg/ bale.\nLoose bales packing)" — SPLIT 2 paragraphs →
        str.replace không span paragraph được → text vẫn hardcode.

    Hàm này detect cell có paragraph bắt đầu bằng "(35kg/" hoặc "(35 kg/" và
    rewrite về single paragraph "({packing_desc})".
    """
    fixed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras = list(cell.paragraphs)
                for i, p in enumerate(paras):
                    text = p.text.strip()
                    # Detect hardcoded packing opener
                    if text.startswith('(35kg/') or text.startswith('(35 kg/'):
                        # Rewrite this paragraph
                        for r in list(p.runs):
                            r._element.getparent().remove(r._element)
                        new_run = p.add_run('({packing_desc})')
                        # Xóa các paragraph kế tiếp thuộc cùng packing block (kết thúc bằng ")")
                        for j in range(i + 1, len(paras)):
                            nxt = paras[j]
                            nxt_text = nxt.text.strip()
                            if not nxt_text:
                                continue
                            if nxt_text.endswith(')') and ('pack' in nxt_text.lower() or 'pallet' in nxt_text.lower() or 'bale' in nxt_text.lower()):
                                nxt._element.getparent().remove(nxt._element)
                                break
                            else:
                                break
                        fixed += 1
                        break
    return fixed


def _inject_extra_terms(doc):
    """
    Tìm paragraph chứa {payment_extra} (SC) hoặc {payment} (PI — trong table cell)
    và insert AFTER nó một paragraph mới:
        {#has_extra_terms}Other Conditions: {extra_terms}{/has_extra_terms}

    docxtemplater paragraphLoop=true sẽ ẩn paragraph khi has_extra_terms=false.
    """
    target = None
    # Pass 1: tìm {payment_extra} trong body paragraphs (SC)
    for p in doc.paragraphs:
        if '{payment_extra}' in (p.text or ''):
            target = p
            break
    # Pass 2: tìm {payment} trong toàn bộ paragraphs (kể cả table cell — PI)
    if target is None:
        for p in _iter_paragraphs(doc):
            text = p.text or ''
            if '{payment}' in text and 'Ben' not in text and 'Bank' not in text:
                target = p
                break

    if target is None:
        return False

    new_text = '{#has_extra_terms}Other Conditions: {extra_terms}{/has_extra_terms}'
    new_para = target.insert_paragraph_before(new_text)
    target._element.addnext(new_para._element)
    return True


def _fix_signature_cell_merge(doc):
    """
    PI templates có signature cell với "HUY ANH RUBBER COMPANY" + "LIMITED" tách
    2 paragraphs → render 2 dòng thay vì 1.
    Merge thành 1 paragraph "HUY ANH RUBBER COMPANY LIMITED".
    """
    fixed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras = list(cell.paragraphs)
                for i, p in enumerate(paras):
                    text = (p.text or '').rstrip()
                    # Detect "HUY ANH RUBBER COMPANY" (without LIMITED) at end of text
                    if text.endswith('HUY ANH RUBBER COMPANY') and i + 1 < len(paras):
                        next_p = paras[i + 1]
                        next_text = (next_p.text or '').strip()
                        if next_text == 'LIMITED':
                            # Append " LIMITED" to current para's first run
                            if p.runs:
                                # Rewrite first run text
                                p.runs[0].text = p.runs[0].text.rstrip() + ' LIMITED'
                                # Remove other runs
                                for r in p.runs[1:]:
                                    r._element.getparent().remove(r._element)
                            else:
                                p.add_run('HUY ANH RUBBER COMPANY LIMITED')
                            # Remove next paragraph (was "LIMITED")
                            next_p._element.getparent().remove(next_p._element)
                            fixed += 1
                            break
    return fixed


def _clean_buyer_address_spacing(doc):
    """
    "ADDRESS      : 295, YANGYEON-RO..." có nhiều space giữa từ ADDRESS và dấu :
    (alignment manual gốc). Clean lại thành "ADDRESS: ...".
    Chỉ áp dụng cho block THE BUYER (không chạm ADDRESS: của bank).
    """
    import re
    fixed = 0
    for p in doc.paragraphs:
        text = p.text or ''
        # Match "ADDRESS<multiple_spaces>:" — but NOT preceded by "BANK" context
        if re.match(r'^ADDRESS\s{2,}:', text):
            if p.runs:
                # Replace pattern in first run
                first = p.runs[0]
                new_text = re.sub(r'^ADDRESS\s+:', 'ADDRESS:', first.text)
                if new_text != first.text:
                    first.text = new_text
                    fixed += 1
    return fixed


def _normalize_last_page(doc):
    """
    Cleanup format các paragraph ở section cuối HĐ (Force Majeure / Arbitration /
    Signature line). Các source DOCX của HA có nhiều chỗ:
      - Para 9.1, 9.2, 9.3 có indent left/right tự đặt (lệch so với 6/7/8)
      - Para 9.3 có leading tab character (extra indent)
      - "FOR THE BUYER ... FOR THE SELLER" có firstLine indent (đẩy text lệch)
      - Empty paragraphs có whitespace thừa
    Hàm này strip leading whitespace + remove indent left/right cho các section
    Force Majeure (9.x), Arbitration (10.x), và signature line.
    """
    from docx.oxml.ns import qn

    def _remove_indents(p, keep_first_line=False):
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            return
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            return
        # Remove left/right/end indent attributes
        for attr in (qn('w:left'), qn('w:right'), qn('w:start'), qn('w:end')):
            if ind.get(attr) is not None:
                del ind.attrib[attr]
        if not keep_first_line:
            for attr in (qn('w:firstLine'), qn('w:hanging')):
                if ind.get(attr) is not None:
                    del ind.attrib[attr]
        # Nếu ind empty thì remove luôn
        if not ind.attrib:
            pPr.remove(ind)

    def _strip_leading_ws(p):
        if not p.runs:
            return
        first = p.runs[0]
        if first.text and first.text != first.text.lstrip():
            first.text = first.text.lstrip()

    changed = 0
    all_paras = doc.paragraphs
    for i, p in enumerate(all_paras):
        text = (p.text or '').strip()
        upper = text.upper()
        # Section 9.x (Force Majeure body) — strip leading WS + remove all indents
        if (text.startswith('9.1 ') or text.startswith('9.2 ') or text.startswith('9.3 ')
            or text.startswith('9.1.') or text.startswith('9.2.') or text.startswith('9.3.')):
            _strip_leading_ws(p)
            _remove_indents(p, keep_first_line=False)
            changed += 1
        # Section 10 (Arbitration body) — keep firstLine indent (paragraph intro), strip other
        elif text.startswith('Arbitration, if any'):
            _strip_leading_ws(p)
            _remove_indents(p, keep_first_line=True)
            changed += 1
        # Signature line "FOR THE BUYER ... FOR THE SELLER" — remove firstLine indent
        elif 'FOR THE BUYER' in upper and 'FOR THE SELLER' in upper:
            _strip_leading_ws(p)
            _remove_indents(p, keep_first_line=False)
            changed += 1
        # Empty paragraphs ở cuối — strip whitespace thừa
        elif not text and p.runs:
            for r in p.runs:
                if r.text and r.text.strip() == '':
                    r.text = ''
    return changed


def build_template(src_path, out_path, replacements, label):
    print(f"\n── {label} ──")
    print(f"  src: {src_path.relative_to(ROOT)}")
    doc = Document(str(src_path))
    changed = 0
    matched_keys = set()
    for p in _iter_paragraphs(doc):
        before = p.text
        if _replace_in_paragraph(p, replacements):
            changed += 1
            for find, repl in replacements:
                if find in before:
                    matched_keys.add(repl)
    # Inject placeholder điều khoản bổ sung sau section Payment
    if _inject_extra_terms(doc):
        matched_keys.add('{#has_extra_terms}...{/has_extra_terms}')
    # Fix PI_FOB hardcoded packing description (split across paragraphs)
    if _fix_pi_packing_in_cells(doc):
        matched_keys.add('[pi_packing_cell_fixed]')
    # Đồng bộ head letter (title sz=44, No/Date sz=28 centered, SELLER block sz=26)
    head_changed = _normalize_head_letter(doc)
    if head_changed:
        matched_keys.add(f'[head_letter normalized: {head_changed} paragraphs]')
    # Cleanup format last page (Force Majeure / Arbitration / Signature)
    last_changed = _normalize_last_page(doc)
    if last_changed:
        matched_keys.add(f'[last_page normalized: {last_changed} paragraphs]')
    # Merge "HUY ANH RUBBER COMPANY" + "LIMITED" 1 line trong PI signature
    sig_fixed = _fix_signature_cell_merge(doc)
    if sig_fixed:
        matched_keys.add(f'[signature_cell merged: {sig_fixed}]')
    # Clean "ADDRESS      :" extra spaces (THE BUYER block)
    addr_fixed = _clean_buyer_address_spacing(doc)
    if addr_fixed:
        matched_keys.add(f'[buyer_address spacing cleaned: {addr_fixed}]')
    doc.save(str(out_path))
    # Copy ra /public/contract-templates/ để Vite serve cho frontend fetch
    public_path = PUBLIC_DIR / out_path.name
    public_path.write_bytes(out_path.read_bytes())
    print(f"  out: {out_path.relative_to(ROOT)}  ({changed} paragraphs modified)")
    print(f"  + copy → {public_path.relative_to(ROOT)}")
    print(f"  placeholders inserted: {sorted(matched_keys)}")
    return changed, matched_keys


def main():
    pairs = [
        (ROOT / "SC + PI ( CIF)" / "SC- Yoong Do engineering Co.,Ltd 01.2026.docx",
         OUT_DIR / "template_SC_CIF.docx", REPL_SC_CIF, "SC CIF (Yoong Do)"),
        (ROOT / "SC + PI ( CIF)" / "PI- Yoong Do engineering Co.,Ltd 01.2026.docx",
         OUT_DIR / "template_PI_CIF.docx", REPL_PI_CIF, "PI CIF (Yoong Do)"),
        (ROOT / "SC + PI (FOB)" / "SC-APOLLO 01.2026 1.docx",
         OUT_DIR / "template_SC_FOB.docx", REPL_SC_FOB, "SC FOB (Apollo)"),
        (ROOT / "SC + PI (FOB)" / "PI- APOLLO 01.2026 1.docx",
         OUT_DIR / "template_PI_FOB.docx", REPL_PI_FOB, "PI FOB (Apollo)"),
    ]
    total = 0
    for src, out, repl, label in pairs:
        if not src.exists():
            print(f"!!! KHÔNG TÌM THẤY: {src}")
            continue
        c, _ = build_template(src, out, repl, label)
        total += c
    print(f"\n✓ Tổng: {total} paragraphs đã modify trong 4 file")


if __name__ == "__main__":
    main()
