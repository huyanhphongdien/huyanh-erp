"""
Generate file DOCX huong dan test luong "Chay dau ra" tren ERP.

Chay: python docs/generate_b2b_chay_dau_ra_guide.py
Output: docs/HUONG_DAN_TEST_B2B_CHAY_DAU_RA.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
PRIMARY = RGBColor(0x1B, 0x4D, 0x3E)
ACCENT_RED = RGBColor(0xDC, 0x26, 0x26)
ACCENT_ORANGE = RGBColor(0xD9, 0x77, 0x06)
ACCENT_BLUE = RGBColor(0x1D, 0x4E, 0xD8)
ACCENT_GREEN = RGBColor(0x16, 0xA3, 0x4A)
ACCENT_CYAN = RGBColor(0x0E, 0xA5, 0xE9)
GRAY = RGBColor(0x6B, 0x72, 0x80)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    if color:
        run.font.color.rgb = color
    elif level == 1:
        run.font.color.rgb = PRIMARY
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_callout(doc, title, body, color_hex='FEF3C7', border_color=ACCENT_ORANGE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_hex)
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    run_t = p_title.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = border_color
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_before = Pt(2)
    run_b = p_body.add_run(body)
    run_b.font.size = Pt(10)
    doc.add_paragraph()


def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


def add_numbered(doc, text):
    return doc.add_paragraph(text, style='List Number')


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1B4D3E')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, 'F9FAFB')
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F3F4F6')
    p = cell.paragraphs[0]
    for i, line in enumerate(code.split('\n')):
        if i > 0:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_cover(doc):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(72)
    run = p_title.add_run('HƯỚNG DẪN TEST LUỒNG "CHẠY ĐẦU RA"')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(8)
    run = p_sub.add_run('B2B Deal lifecycle 10 bước — DRC after production')
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(40)
    run = p_meta.add_run('Đối tượng: Minh LD + Phú LV + Đại lý partner')
    run.bold = True
    run.font.size = Pt(12)

    p_phase = doc.add_paragraph()
    p_phase.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_phase.paragraph_format.space_before = Pt(4)
    run = p_phase.add_run('Sau Phase 1-3 fixes (SQL audit verified 2026-05-23)')
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = ACCENT_GREEN

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(4)
    run = p_date.add_run('Cập nhật: 2026-05-23')
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(80)
    run = p_foot.add_run('HUY ANH RUBBER COMPANY LIMITED')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = PRIMARY

    p_foot2 = doc.add_paragraph()
    p_foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_foot2.add_run('ERP System · huyanhrubber.vn')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# BUILD GUIDE
# ═══════════════════════════════════════════════════════════════════════════════

def build_guide():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    add_cover(doc)

    # ═══ Tổng quan ═══
    add_heading(doc, 'Tổng quan luồng "Chạy đầu ra"', level=1)
    add_para(doc, '"Chạy đầu ra" (DRC after production) là loại deal mua mủ tươi từ đại lý — KHÁC với "Mua đứt":')

    add_table(doc,
        headers=['So sánh', 'Mua đứt', 'Chạy đầu ra'],
        rows=[
            ['Giá chốt', 'Giá USD/MT ngay lúc chốt', 'Giá tạm tính, chốt sau SX'],
            ['DRC dùng để tính giá', 'Sample DRC lúc nhận hàng', 'Actual DRC sau khi SX TP'],
            ['Số bước lifecycle', '4 bước (Chốt→Duyệt→Nhập kho→Thanh toán)', '10 bước (đầy đủ SX)'],
            ['Khi nào dispute', 'Khi sample sai > 3% với KH yêu cầu', 'Khi actual_drc lệch sample_drc > 3%'],
            ['Use case', 'Mủ chuẩn, KH biết DRC', 'Mủ tạp, DRC không chắc'],
        ],
        col_widths=[5, 5, 6]
    )

    add_heading(doc, '10 bước lifecycle Chạy đầu ra', level=2)
    add_code_block(doc, '''1.  Đã cân          — Weighbridge IN (app cân tại TL/PD)
2.  Đã nhập kho     — Auto stock_in từ ticket (trigger)
3.  QC sample DRC   — Factory bấm "Nhập sample DRC"      ← UI mới Phase 2
4.  BGĐ duyệt       — Status: processing → accepted
5.  Tạm ứng         — Advance + Partner ack
6.  Bắt đầu SX      — Factory bấm "Bắt đầu sản xuất"     ← UI mới Phase 2
7.  Ra TP           — Factory bấm "Hoàn tất sản xuất"    ← UI mới Phase 2
8.  QC final        — Auto compute actual_drc + dispute check
9.  Quyết toán      — Factory tạo settlement
10. Thanh toán      — Paid → status=settled''')

    # ═══ Pre-conditions ═══
    add_heading(doc, 'Pre-conditions', level=1)

    add_para(doc, 'Trước khi bắt đầu test, đảm bảo:', bold=True)
    add_numbered(doc, 'Anh login với tài khoản Minh LD (minhld@huyanhrubber.com) — có toàn quyền')
    add_numbered(doc, 'Có 1 đại lý đã verified — vd Goutam Enterprises, Apollo, hoặc partner test')
    add_numbered(doc, 'TEST MODE email đang ON (chỉ Minh nhận mail) — em đã verify')
    add_numbered(doc, 'Sample deal sẵn để test: DL2605-SI2O (110 tấn mủ tạp, 24,000 đ/kg) — đã ở ERP')

    add_callout(doc,
        '💡 SQL verify trước khi test',
        'Em đã chạy SQL audit 2026-05-23 confirm: schema có đủ 8/8 cột lifecycle, '
        'trigger trg_drc_variance_dispute active, b2b_notifications accept 3 type mới. '
        'Production-ready để test.',
        color_hex='DCFCE7', border_color=ACCENT_GREEN
    )

    # ═══ Truy cập màn hình ═══
    add_heading(doc, 'Truy cập màn hình Deal', level=1)

    add_para(doc, 'Sidebar bên trái → group "B2B THU MUA" → "Deals"', bold=True)

    add_table(doc,
        headers=['Menu', 'URL', 'Mục đích'],
        rows=[
            ['Chat Đại lý', '/b2b/chat', 'Tạo Deal mới từ booking đại lý'],
            ['Nhu cầu mua', '/b2b/demands', 'Đăng nhu cầu mua chiêu mộ đại lý'],
            ['Đại lý', '/b2b/partners', 'Quản lý đại lý'],
            ['Deals', '/b2b/deals', '★ Danh sách deal — click row → detail'],
            ['Đấu giá', '/b2b/auctions', 'Tạo đấu giá'],
            ['Nhập kho', '/b2b/rubber-intake', 'Wizard intake (outright/walkin/production)'],
            ['Công nợ', '/b2b/ledger', 'Sổ cái đại lý'],
            ['Quyết toán', '/b2b/settlements', 'Phiếu quyết toán'],
        ],
        col_widths=[3.5, 4, 8]
    )

    add_para(doc, 'Bước test này tập trung vào /b2b/deals/{id} — Drawer/Page Detail của Deal.', italic=True)

    doc.add_page_break()

    # ═══ TC-1 ═══
    add_heading(doc, 'TC-1 (Bước 1-2): Cân + Nhập kho', level=1, color=ACCENT_CYAN)
    add_para(doc, 'Mục đích: tạo phiếu cân IN + auto stock-in vào kho NVL', italic=True)

    add_heading(doc, 'Action', level=2)
    add_numbered(doc, 'Mở app cân (Tân Lâm hoặc Phong Điền)')
    add_numbered(doc, 'Chọn deal DL2605-SI2O từ dropdown')
    add_numbered(doc, 'Cân lần 1 (gross) → save')
    add_numbered(doc, 'Cân lần 2 (tare) → tính NET')
    add_numbered(doc, 'Bấm Submit ticket → auto-trigger stock_in')

    add_heading(doc, 'Verify trong ERP', level=2)
    add_bullet(doc, 'Mở /b2b/deals/{id} → tab "Nhập kho (1)" — badge số stock_in')
    add_bullet(doc, 'Thấy 1 row stock_in_order status=confirmed')
    add_bullet(doc, 'Tab "Sản xuất" — timeline bước 1 "Đã cân" + bước 2 "Đã nhập kho" → done (xanh)')

    add_heading(doc, 'SQL verify', level=2)
    add_code_block(doc, '''SELECT deal_number, stock_in_count, status
FROM b2b.deals WHERE deal_number = 'DL2605-SI2O';
-- Expected: stock_in_count > 0

SELECT code, ticket_type, status, completed_at
FROM weighbridge_tickets
WHERE deal_id = '<deal_id>';
-- Expected: 1 row ticket_type='in' status='completed' ''')

    # ═══ TC-2 ═══
    add_heading(doc, 'TC-2 (Bước 3): Nhập Sample DRC ⭐ UI MỚI', level=1, color=ACCENT_CYAN)
    add_para(doc, 'Mục đích: QC đo mẫu mủ vừa nhập kho, ghi sample DRC để BGĐ có cơ sở duyệt', italic=True)

    add_heading(doc, 'Action', level=2)
    add_numbered(doc, 'Vào /b2b/deals/{deal_id} (vd /b2b/deals/ad548d48-...)')
    add_numbered(doc, 'Click tab "Sản xuất" (icon tên lửa 🚀) — chỉ hiện khi purchase_type=drc_after_production')
    add_numbered(doc, 'Thấy Timeline 10 stages + section button bên dưới')
    add_numbered(doc, 'Click button XANH NHẠT "🧪 Nhập sample DRC"')
    add_numbered(doc, 'Modal mở:\n  - Alert info: "QC lấy mẫu mủ ngay sau khi nhập kho..."\n  - Input "Sample DRC (%)" autofocus\n  - Nhập 35.5 → click "Lưu sample DRC"')

    add_heading(doc, 'Verify', level=2)
    add_bullet(doc, 'Toast xanh "Đã ghi sample DRC = 35.5%"')
    add_bullet(doc, 'Timeline cập nhật: stage "QC sample DRC" → done')
    add_bullet(doc, 'Button "Nhập sample DRC" biến mất (đã xong)')
    add_bullet(doc, 'Email Minh LD nhận notification "🧪 Đã ghi nhận Sample DRC"')

    add_heading(doc, 'Edge cases (em đã code)', level=2)
    add_table(doc,
        headers=['Input', 'Behavior'],
        rows=[
            ['DRC = 0 hoặc 101', 'Modal validate fail "DRC ∈ (0, 100]"'],
            ['stock_in_count = 0', 'Service throw "Phải nhập kho ít nhất 1 batch trước"'],
            ['status != processing', 'Service throw "Phải ở status=processing"'],
            ['purchase_type != drc_after_production', 'Service throw "Chỉ áp dụng cho Chạy đầu ra"'],
        ],
        col_widths=[6, 10]
    )

    # ═══ TC-3 ═══
    add_heading(doc, 'TC-3 (Bước 4): BGĐ duyệt Deal', level=1, color=ACCENT_CYAN)
    add_para(doc, 'Mục đích: Trung/Huy/Minh LD duyệt → status: processing → accepted', italic=True)

    add_heading(doc, 'Action', level=2)
    add_numbered(doc, 'Trên DealDetailPage, ở header section thấy nút "Duyệt Deal" (chỉ khi status=processing + đủ điều kiện)')
    add_numbered(doc, 'HOẶC vào /b2b/chat → tìm DealCard trong room → click "Duyệt Deal" (xanh)')
    add_numbered(doc, 'Modal "Duyệt Deal" mở:\n  - Hiển thị info deal\n  - Optional: final_price + notes\n  - Click "Duyệt"')

    add_heading(doc, 'Verify', level=2)
    add_bullet(doc, 'Status badge đổi: Processing (cam) → Đã duyệt (xanh)')
    add_bullet(doc, 'Timeline stage "BGĐ duyệt" → done')
    add_bullet(doc, 'Tab "Sản xuất" — bước "Tạm ứng" → current')

    add_heading(doc, 'SQL verify audit log', level=2)
    add_code_block(doc, '''SELECT op, changed_fields,
       old_data->>'status' AS old_status,
       new_data->>'status' AS new_status,
       changed_at
FROM b2b.deal_audit_log
WHERE deal_id = '<deal_id>'
ORDER BY changed_at DESC LIMIT 3;
-- Expected: 1 row op='UPDATE' old=processing new=accepted''')

    add_callout(doc,
        '⚠ Schema gotcha',
        'Table audit log ở b2b.deal_audit_log (KHÔNG phải public.deal_audit_log). '
        'Trigger trg_deal_audit tự fire khi UPDATE deal.',
        color_hex='FEF3C7', border_color=ACCENT_ORANGE
    )

    # ═══ TC-4 ═══
    add_heading(doc, 'TC-4 (Bước 5): Tạm ứng', level=1, color=ACCENT_CYAN)
    add_para(doc, 'Mục đích: Factory tạm ứng tiền cho đại lý đi mua nguyên liệu', italic=True)

    add_heading(doc, 'Action ERP', level=2)
    add_numbered(doc, 'DealDetailPage tab "Tạm ứng" (icon ví 💰)')
    add_numbered(doc, 'HOẶC DealCard chat → click "Ứng thêm" (deal status=accepted)')
    add_numbered(doc, 'Modal mở:\n  - Amount: 1,000,000,000 đ\n  - Purpose: "Mua nguyên liệu mủ tươi"\n  - Payment method: cash / bank_transfer\n  - Submit')

    add_heading(doc, 'Action Partner Portal', level=2)
    add_numbered(doc, 'Đại lý vào b2b.huyanhrubber.vn (Portal)')
    add_numbered(doc, 'Tab Tài chính → Tạm ứng')
    add_numbered(doc, 'Click "Xác nhận đã nhận" → status: pending → acknowledged')

    add_heading(doc, 'Verify', level=2)
    add_bullet(doc, 'b2b_advances 1 row status=acknowledged (sau ack)')
    add_bullet(doc, 'DealCard ERP hiện "Đã ứng 1 lần"')
    add_bullet(doc, 'Timeline stage "Tạm ứng" → done')

    # ═══ TC-5 ═══
    add_heading(doc, 'TC-5 (Bước 6): Bắt đầu sản xuất ⭐ UI MỚI', level=1, color=ACCENT_CYAN)
    add_para(doc, 'Mục đích: nhà máy chính thức start SX lô hàng đại lý', italic=True)

    add_heading(doc, 'Action', level=2)
    add_numbered(doc, 'Vào /b2b/deals/{id} → tab "Sản xuất"')
    add_numbered(doc, 'Sau khi sample_drc + accepted + advance ack → thấy button CAM "🏭 Bắt đầu sản xuất"')
    add_numbered(doc, 'Click → Modal confirm hiện:\n  - Deal number\n  - NL đầu vào (kg)\n  - Sample DRC đã chốt\n  - Cảnh báo "Sau khi bấm: production_started_at = NOW()"')
    add_numbered(doc, 'Click "Bắt đầu SX"')

    add_heading(doc, 'Verify', level=2)
    add_bullet(doc, 'Toast "Đã bắt đầu sản xuất Deal DL2605-SI2O"')
    add_bullet(doc, 'DB b2b_deals.production_started_at = NOW()')
    add_bullet(doc, 'Timeline stage "Bắt đầu sản xuất" → done, "Ra thành phẩm" → current')
    add_bullet(doc, 'Email Minh nhận notification "🏭 Nhà máy đã bắt đầu sản xuất"')

    add_heading(doc, 'Edge cases', level=2)
    add_table(doc,
        headers=['Trạng thái sai', 'Service phản hồi'],
        rows=[
            ['status != accepted', '"Phải ở status=accepted"'],
            ['production_started_at != null (đã start)', '"Đã start từ trước"'],
            ['No advance acknowledged', '"Cần ack advance trước"'],
            ['sample_drc null', '(button không hiện)'],
        ],
        col_widths=[6, 10]
    )

    # ═══ TC-6 ═══
    add_heading(doc, 'TC-6 (Bước 7-8): Hoàn tất SX + QC final ⭐ UI MỚI', level=1, color=ACCENT_RED)
    add_para(doc, 'Mục đích: nhập KL thành phẩm sau SX → ERP tự compute actual_drc + giá cuối + raise dispute nếu lệch sample > 3%.', italic=True)
    add_para(doc, '⚠ ĐÂY LÀ BƯỚC QUAN TRỌNG NHẤT — quyết định giá cuối thanh toán cho đại lý.', bold=True, color=ACCENT_RED)

    add_heading(doc, 'Scenario A — Variance OK (≤ 3%)', level=2)
    add_numbered(doc, 'Tab "Sản xuất" — sau khi production_started_at có → thấy button XANH LÁ "✅ Hoàn tất sản xuất + QC final"')
    add_numbered(doc, 'Click → Modal mở:\n  - Info: NL=110,000 kg, đơn giá 24,000 đ/kg, sample DRC=35.5%\n  - Input "Khối lượng thành phẩm (kg)" autofocus\n  - Nhập 39,050')
    add_numbered(doc, 'Preview LIVE hiện:\n  - Actual DRC = 39,050 / 110,000 × 100 = 35.50%\n  - Variance vs Sample = 0.00% (OK)\n  - Giá cuối = 110,000 × 35.5% × 24,000 = 937,200,000 đ\n  - Alert xanh "Preview giá cuối"')
    add_numbered(doc, 'Click "Chốt KL thành phẩm + Auto-compute giá cuối"')

    add_heading(doc, 'Verify Scenario A', level=3)
    add_bullet(doc, 'Toast "SX xong: actual DRC=35.50%, giá cuối=937,200,000đ"')
    add_bullet(doc, 'DB updated: actual_drc=35.50, finished_product_kg=39050, final_value=937200000')
    add_bullet(doc, 'Timeline stage "Ra thành phẩm" + "QC final" → done')
    add_bullet(doc, 'KHÔNG có dispute mới (variance < 3%)')
    add_bullet(doc, 'Email Minh "✅ Sản xuất xong — giá cuối chốt"')

    add_heading(doc, 'Scenario B — Variance > 3% (auto-dispute)', level=2)
    add_para(doc, 'Repeat TC-6 với finished_product_kg = 30,800 kg:', bold=True)
    add_bullet(doc, 'Actual DRC = 30,800 / 110,000 = 28.00%')
    add_bullet(doc, 'Variance vs sample 35.5% = 7.50% (> 3% → DISPUTE!)')
    add_bullet(doc, 'Giá cuối giảm = 110,000 × 28% × 24,000 = 739,200,000 đ')
    add_bullet(doc, 'Preview Alert VÀNG (warning) "Variance > 3% → auto-raise dispute"')

    add_heading(doc, 'Verify Scenario B', level=3)
    add_bullet(doc, 'Toast WARNING "Variance > 3% → auto-raise dispute!"')
    add_bullet(doc, 'b2b_drc_disputes có 1 row mới status=open (trigger trg_drc_variance_dispute fire)')
    add_bullet(doc, 'Email Minh "⚠️ Sản xuất xong — variance DRC > 3% — vui lòng review"')

    add_heading(doc, 'SQL verify dispute', level=3)
    add_code_block(doc, '''SELECT dispute_number, expected_drc, actual_drc,
       drc_variance, reason, status, raised_at
FROM b2b.drc_disputes
WHERE deal_id = '<deal_id>'
ORDER BY raised_at DESC LIMIT 1;
-- Expected (Scenario B): 1 row status=open, drc_variance = 7.5%''')

    add_callout(doc,
        '🔬 Trigger name thực',
        'Trigger fire là trg_drc_variance_dispute (KHÔNG phải tên "P16" trong code comments). '
        'Trigger check abs(actual_drc - sample_drc) > 3% thì auto INSERT vào b2b.drc_disputes.',
        color_hex='E0F2FE', border_color=ACCENT_BLUE
    )

    # ═══ TC-7 ═══
    add_heading(doc, 'TC-7 (Bước 9-10): Quyết toán + Thanh toán', level=1, color=ACCENT_CYAN)
    add_para(doc, 'Mục đích: Factory tạo settlement + payment final', italic=True)

    add_heading(doc, 'Action', level=2)
    add_numbered(doc, 'DealCard chat hoặc DealDetailPage → click "Tạo quyết toán" (purple)')
    add_numbered(doc, 'Modal auto-fill final_value đã tính ở TC-6')
    add_numbered(doc, 'Confirm → status: draft → pending_approval')
    add_numbered(doc, 'Manager approve → status: approved')
    add_numbered(doc, 'Mark paid khi thanh toán xong → status: paid')

    add_heading(doc, 'Verify', level=2)
    add_bullet(doc, 'b2b_settlements 1 row status=paid')
    add_bullet(doc, 'b2b_deals.status = settled')
    add_bullet(doc, 'Timeline stage "Quyết toán" + "Thanh toán" → done (full xanh 10/10)')

    # ═══ Test Matrix ═══
    add_heading(doc, 'Test Matrix tóm tắt', level=1)

    add_table(doc,
        headers=['TC', 'Stage', 'Trước', 'Sau', 'UI Phase 2?', 'Notify Partner?'],
        rows=[
            ['1', 'Cân + Nhập kho', 'new', 'stock_in_count>0', '—', '—'],
            ['2', 'QC sample DRC', 'stock_in>0', 'sample_drc set', '✅', '✅'],
            ['3', 'BGĐ duyệt', 'processing', 'accepted', '— (sẵn)', '—'],
            ['4', 'Tạm ứng', 'accepted', 'advance ack', '— (sẵn)', '—'],
            ['5', 'Start SX', 'accepted+adv', 'production_started_at', '✅', '✅'],
            ['6', 'Finish SX + QC', 'production_started', 'actual_drc + final_value', '✅', '✅ (+dispute)'],
            ['7', 'Settlement + Paid', 'actual_drc set', 'settled', '— (sẵn)', '—'],
        ],
        col_widths=[1.5, 3, 3, 3.5, 2.5, 2.5]
    )

    # ═══ Common issues ═══
    add_heading(doc, 'Sự cố thường gặp', level=1)

    add_heading(doc, 'Sự cố 1: Tab "Sản xuất" KHÔNG hiện', level=2, color=ACCENT_RED)
    add_para(doc, 'Nguyên nhân: purchase_type của deal != "drc_after_production"', italic=True)
    add_para(doc, 'Verify:', bold=True)
    add_code_block(doc, '''SELECT deal_type, purchase_type FROM b2b.deals
WHERE deal_number = 'DL2605-SI2O';
-- Expected: processing | drc_after_production''')
    add_para(doc, 'Fix: Phase 1 code đã sync deal_type→purchase_type. Nếu deal cũ chưa sync, run migration b2b_sync_purchase_type_v16.sql.')

    add_heading(doc, 'Sự cố 2: Button "Nhập sample DRC" KHÔNG hiện', level=2, color=ACCENT_RED)
    add_para(doc, 'Nguyên nhân: stock_in_count = 0 (chưa cân + nhập kho TC-1)', italic=True)
    add_para(doc, 'Fix: hoàn thành TC-1 trước → cân + auto stock_in.')

    add_heading(doc, 'Sự cố 3: Button "Bắt đầu sản xuất" KHÔNG hiện', level=2, color=ACCENT_RED)
    add_para(doc, 'Nguyên nhân: Một trong các điều kiện chưa đủ:', italic=True)
    add_bullet(doc, 'status != accepted (chưa BGĐ duyệt — TC-3)')
    add_bullet(doc, 'sample_drc = null (chưa TC-2)')
    add_bullet(doc, 'Chưa có advance acknowledged (chưa hoặc đại lý chưa ack TC-4)')

    add_heading(doc, 'Sự cố 4: Hoàn tất SX báo lỗi RLS permission denied', level=2, color=ACCENT_RED)
    add_para(doc, 'Nguyên nhân: User login không có quyền UPDATE b2b.deals', italic=True)
    add_para(doc, 'Fix: Login với role employee + có dept thuộc factory.')

    add_heading(doc, 'Sự cố 5: Email notification không nhận', level=2, color=ACCENT_RED)
    add_para(doc, 'Nguyên nhân khả dĩ:', italic=True)
    add_bullet(doc, 'Email tại Minh — TEST MODE đang ON nên chỉ Minh nhận (không phải bug)')
    add_bullet(doc, 'Check Spam folder')
    add_bullet(doc, 'Edge function b2b-deal-notify chưa deploy: npx supabase functions deploy b2b-deal-notify')

    # ═══ Acceptance criteria ═══
    add_heading(doc, 'Acceptance criteria — Test PASS khi:', level=1)
    add_bullet(doc, 'Cả 7 TC chạy hết không có lỗi')
    add_bullet(doc, 'Timeline progression đúng 10 stages')
    add_bullet(doc, 'Final value match công thức: quantity_kg × actual_drc% × unit_price')
    add_bullet(doc, 'Dispute auto-raise khi variance > 3%')
    add_bullet(doc, 'Partner notifications 3 events (sample/start/finish)')
    add_bullet(doc, 'Settlement final_value = deal.final_value')
    add_bullet(doc, 'Status flow đúng: processing → accepted → settled')

    # ═══ Test log ═══
    add_heading(doc, 'Test execution log', level=1)

    add_table(doc,
        headers=['Ngày', 'Tester', 'TC', 'Kết quả', 'Ghi chú'],
        rows=[
            ['2026-05-23', 'Claude SQL audit', 'Schema', '✅ 15/17', 'Production-ready, 2 typo guide đã fix'],
            ['', '(chưa test)', 'TC-1', '—', ''],
            ['', '(chưa test)', 'TC-2', '—', ''],
            ['', '(chưa test)', 'TC-3', '—', ''],
            ['', '(chưa test)', 'TC-4', '—', ''],
            ['', '(chưa test)', 'TC-5', '—', ''],
            ['', '(chưa test)', 'TC-6', '—', ''],
            ['', '(chưa test)', 'TC-7', '—', ''],
        ],
        col_widths=[3, 4, 2, 2, 5]
    )

    # ═══ Footer ═══
    add_callout(doc,
        '📞 Liên hệ khi có vấn đề',
        'UI lỗi / button không hiện / SQL fail → báo Minh LD. '
        'Edge case nghiệp vụ chưa cover (vd multi-batch, multi-pool) → thảo luận BGĐ trước khi production.',
        color_hex='E0F2FE', border_color=ACCENT_BLUE
    )

    return doc


def main():
    doc = build_guide()
    path = ROOT / 'HUONG_DAN_TEST_B2B_CHAY_DAU_RA.docx'
    doc.save(str(path))
    print(f'OK Generated: {path.name} ({path.stat().st_size // 1024} KB)')


if __name__ == '__main__':
    main()
