"""
Generate file DOCX: HUONG_DAN_DOCS_TEMPLATE_HOP_DONG.docx

Hướng dẫn Docs (Nhung + Phương Anh) đổi 4 file template HĐ từ
highlight vàng sang placeholder {{...}} để ERP auto-fill.

Chạy: python docs/generate_docs_template_guide.py
Output: docs/HUONG_DAN_DOCS_TEMPLATE_HOP_DONG.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
PRIMARY = RGBColor(0x1B, 0x4D, 0x3E)
ACCENT_RED = RGBColor(0xDC, 0x26, 0x26)
ACCENT_ORANGE = RGBColor(0xD9, 0x77, 0x06)
ACCENT_BLUE = RGBColor(0x1D, 0x4E, 0xD8)
ACCENT_GREEN = RGBColor(0x16, 0xA3, 0x4A)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    if color:
        run.font.color.rgb = color
    elif level == 1:
        run.font.color.rgb = PRIMARY
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_callout(doc, title, body, color_hex='FEF3C7', border_color=ACCENT_ORANGE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_hex)
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    run_t = p_title.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = border_color
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_before = Pt(2)
    run_b = p_body.add_run(body)
    run_b.font.size = Pt(10)
    doc.add_paragraph()


def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


def add_numbered(doc, text):
    return doc.add_paragraph(text, style='List Number')


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1B4D3E')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, 'F9FAFB')
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, code, monospace=True):
    """Code block — table 1x1 với background xám nhạt + font monospace."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F3F4F6')
    p = cell.paragraphs[0]
    for i, line in enumerate(code.split('\n')):
        if i > 0:
            p = cell.add_paragraph()
        run = p.add_run(line)
        if monospace:
            run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_cover(doc):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(72)
    run = p_title.add_run('HƯỚNG DẪN ĐỔI TEMPLATE HỢP ĐỒNG')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(8)
    run = p_sub.add_run('Từ Highlight vàng → Placeholder text {{...}}')
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(36)
    run = p_meta.add_run('Đối tượng: Nhung + Phương Anh (Docs)')
    run.bold = True
    run.font.size = Pt(12)

    p_time = doc.add_paragraph()
    p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_time.paragraph_format.space_before = Pt(4)
    run = p_time.add_run('Thời gian thực hiện: ~10 phút (làm 1 lần duy nhất, lợi mãi)')
    run.font.size = Pt(11)
    run.italic = True

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(4)
    run = p_date.add_run('Cập nhật: 2026-05-22')
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(120)
    run = p_foot.add_run('HUY ANH RUBBER COMPANY LIMITED')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = PRIMARY

    p_foot2 = doc.add_paragraph()
    p_foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_foot2.add_run('ERP System · huyanhrubber.vn')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
# BUILD DOC
# ═════════════════════════════════════════════════════════════════════════════

def build_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    add_cover(doc)

    # ─── 0. Tại sao 2 ngoặc thay vì 1 (giải thích Nhung's confusion) ───
    add_heading(doc, 'Tại sao 2 ngoặc {{...}} thay vì 1 ngoặc {...}?', level=1, color=ACCENT_RED)

    add_callout(doc,
        '⚠ Lưu ý quan trọng',
        'Nhung/Phương Anh có thể đã thấy các template CŨ trong máy với '
        '{contract_no}, {contract_date}... (1 ngoặc). Đó là format compose flow CŨ '
        '(deprecated) — ERP cũ tự render TOÀN BỘ HĐ. Workflow đó không còn dùng.',
        color_hex='FEE2E2',
        border_color=ACCENT_RED
    )

    add_para(doc, 'Workflow MỚI (upload flow) — Docs upload file, ERP chỉ fill 6 trường — bắt buộc dùng 2 ngoặc:', bold=True)

    add_table(doc,
        headers=['Format', 'Ý nghĩa', 'Khi nào dùng'],
        rows=[
            ['{contract_no}\n(1 ngoặc)', 'Format CŨ (compose flow deprecated) — ERP cũ tự render full HĐ', 'Đừng dùng cho HĐ mới'],
            ['{{contract_no}}\n(2 ngoặc)', 'Format MỚI (upload flow) — ERP chỉ fill 6 trường', 'Dùng cho mọi HĐ từ giờ'],
        ],
        col_widths=[4, 8, 5]
    )

    add_heading(doc, 'Lý do dùng 2 ngoặc:', level=3)
    add_para(doc, 'Tránh nhầm với các text bình thường có dấu { hoặc } trong HĐ (vd: ghi chú, công thức, ký hiệu). 2 ngoặc gần như không bao giờ xuất hiện ngẫu nhiên trong HĐ thực tế.')

    add_heading(doc, 'Action cho chị Nhung/Phương Anh:', level=3)
    add_numbered(doc, 'Nếu template hiện tại có {contract_no} (1 ngoặc) → đổi thành {{contract_no}} (2 ngoặc) — chỉ cần thêm 1 ngoặc mỗi bên')
    add_numbered(doc, 'Áp dụng tương tự cho 5 token bank (xem bảng bên dưới)')
    add_numbered(doc, 'Các token CŨ khác (vd: {buyer_name}, {grade}, {quantity}...) → XÓA và gõ giá trị THẬT vào per HĐ. Workflow mới Docs custom HĐ theo từng khách, ERP không tự fill các trường này.')

    add_callout(doc,
        '💡 Quick check',
        'Trong Word bấm Ctrl+F → tìm "{" (1 ngoặc) → nếu còn token nào dạng 1 ngoặc '
        'trong template chính, sửa hết thành 2 ngoặc cho 6 token chính, hoặc thay '
        'bằng giá trị thật cho các token khác.',
        color_hex='E0F2FE',
        border_color=ACCENT_BLUE
    )

    doc.add_page_break()

    # ─── 1. Vì sao phải đổi ───
    add_heading(doc, 'Vì sao phải đổi?', level=1)

    add_heading(doc, 'Tình trạng hiện tại', level=2)
    add_para(doc, 'Mỗi HĐ mới, quy trình giữa Docs ↔ Phú đang là:')
    add_numbered(doc, 'Docs: Mở template gốc → copy ra HĐ mới → custom buyer/grade/giá → để HIGHLIGHT VÀNG ở 2 chỗ (số HĐ + bank) → upload ERP')
    add_numbered(doc, 'Phú LV: Download file → mở Word → click vào vùng vàng → gõ số HĐ + bank info → save → upload lại ERP')
    add_numbered(doc, 'Mất ~3 phút/HĐ cho Phú × N HĐ/tuần = nhiều thời gian')

    add_heading(doc, 'Sau khi đổi', level=2)
    add_numbered(doc, 'Docs: Mở template (đã sửa sẵn theo hướng dẫn này) → custom buyer/grade/giá → upload ERP (giữ nguyên các {{...}})')
    add_numbered(doc, 'Phú LV: Vào ERP, gõ số HĐ + chọn ngân hàng từ dropdown → bấm "Auto-fill" → ERP tự fill vào file → bấm Duyệt')
    add_numbered(doc, 'Chỉ ~30 giây/HĐ. Docs không tốn thêm thời gian gì so với hiện tại.')

    # ─── 2. 4 file template ───
    add_heading(doc, '4 file template cần sửa', level=1)
    add_para(doc, 'Đây là file master của Docs (lưu chỗ nào quen — Drive, máy cá nhân, OneDrive...). Có 4 biến thể tùy loại HĐ:')

    add_table(doc,
        headers=['File', 'Khi nào dùng'],
        rows=[
            ['HA_template_SC_FOB.docx', 'Sales Contract (HĐ chính) — Incoterm FOB'],
            ['HA_template_SC_CIF.docx', 'Sales Contract (HĐ chính) — Incoterm CIF'],
            ['HA_template_PI_FOB.docx', 'Proforma Invoice / Commercial Invoice — FOB'],
            ['HA_template_PI_CIF.docx', 'Proforma Invoice / Commercial Invoice — CIF'],
        ],
        col_widths=[6, 10]
    )

    add_callout(doc,
        '💡 Lưu ý',
        'Tên file có thể khác tùy chị đang đặt. Cứ tìm 4 file gốc tương ứng 4 loại HĐ này.',
        color_hex='E0F2FE',
        border_color=ACCENT_BLUE
    )

    # ─── 3. Cách sửa ───
    add_heading(doc, 'Cách sửa — Step by step', level=1)

    add_heading(doc, 'Bước 1: Mở file template gốc trong Word', level=2)
    add_para(doc, 'VD: HA_template_SC_FOB.docx')

    add_heading(doc, 'Bước 2: Tìm các vùng HIGHLIGHT VÀNG (số HĐ + bank info)', level=2)
    add_para(doc, 'Mở file ra, tìm những đoạn đang để trống/có highlight vàng. Có 6 vùng điển hình:')

    add_table(doc,
        headers=['#', 'Vùng', 'Vị trí trong HĐ'],
        rows=[
            ['1', 'Số HĐ', 'Đầu HĐ — "SALES CONTRACT No.: ___"'],
            ['2', 'Tên TK ngân hàng', 'Cuối HĐ — "Account Name: ___"'],
            ['3', 'Số TK', '"Account No.: ___"'],
            ['4', 'Tên đầy đủ ngân hàng', '"Bank: ___"'],
            ['5', 'Địa chỉ ngân hàng', '"Address: ___"'],
            ['6', 'SWIFT code', '"SWIFT: ___"'],
        ],
        col_widths=[1, 5, 10]
    )

    add_heading(doc, 'Bước 3: Xóa highlight vàng, gõ vào đúng vùng đó TOKEN TEXT', level=2)
    add_para(doc, 'Thay vùng trống/highlight bằng các token chính xác sau. CHÚ Ý: CHÍNH XÁC từng ký tự, không dấu cách thừa.')

    add_table(doc,
        headers=['Vùng', 'Gõ vào file'],
        rows=[
            ['Số HĐ', '{{contract_no}}'],
            ['Tên TK ngân hàng', '{{bank_account_name}}'],
            ['Số TK', '{{bank_account_no}}'],
            ['Tên đầy đủ ngân hàng', '{{bank_full_name}}'],
            ['Địa chỉ ngân hàng', '{{bank_address}}'],
            ['SWIFT code', '{{bank_swift}}'],
        ],
        col_widths=[7, 9]
    )

    add_heading(doc, 'Bước 4: Xóa background vàng', level=2)
    add_para(doc, 'Sau khi gõ token vào, có thể vùng đó vẫn còn highlight vàng từ trước. Bôi đen → menu Home → Text Highlight Color → No Color để xóa.')

    add_heading(doc, 'Bước 5: Save file (Ctrl+S)', level=2)
    add_para(doc, 'Quan trọng: Save dưới dạng .docx (không phải .doc).', bold=True, color=ACCENT_RED)

    add_heading(doc, 'Bước 6: Lặp lại Bước 1-5 cho 3 file template còn lại', level=2)

    # ─── 4. Ví dụ Trước / Sau ───
    add_heading(doc, 'Ví dụ Trước / Sau', level=1)

    add_heading(doc, 'TRƯỚC (file template hiện tại)', level=2, color=ACCENT_RED)
    add_code_block(doc, '''SALES CONTRACT
No.: ░░░░░░░░░░░░          ← highlight vàng, trống
Date: 22 May 2026

Buyer: APOLLO TYRES LTD
Address: ...

...

PAYMENT:
Account Name: ░░░░░░░░░░░░░░░░░░░░░░    ← highlight vàng
Account No.:  ░░░░░░░░░░                ← highlight vàng
Bank:         ░░░░░░░░░░░░░░░░░░░░░░    ← highlight vàng
Address:      ░░░░░░░░░░░░░░░░░░░░░░    ← highlight vàng
SWIFT:        ░░░░░░░░                  ← highlight vàng''')

    add_heading(doc, 'SAU (file template sửa xong)', level=2, color=ACCENT_GREEN)
    add_code_block(doc, '''SALES CONTRACT
No.: {{contract_no}}                    ← token, KHÔNG highlight
Date: 22 May 2026

Buyer: APOLLO TYRES LTD
Address: ...

...

PAYMENT:
Account Name: {{bank_account_name}}
Account No.:  {{bank_account_no}}
Bank:         {{bank_full_name}}
Address:      {{bank_address}}
SWIFT:        {{bank_swift}}''')

    # ─── 5. 5 lỗi thường gặp ───
    add_heading(doc, '5 LỖI THƯỜNG GẶP — TRÁNH NGAY', level=1, color=ACCENT_RED)

    add_heading(doc, 'Lỗi 1: Format trong token', level=2, color=ACCENT_RED)
    add_para(doc, 'SAI: Bôi đậm/in nghiêng/màu khác ngay GIỮA token.', bold=True)
    add_code_block(doc, '{{contract_no}}     ← phần "contract" bị bôi đậm')
    add_para(doc, 'ĐÚNG: Toàn bộ token cùng 1 format (cùng font, cùng size, cùng màu).', bold=True, color=ACCENT_GREEN)
    add_callout(doc,
        '💡 Cách test',
        'Bôi đen toàn bộ {{contract_no}} (cả 2 dấu {{ lẫn }}) — nếu hiện B ở thanh format khi click vào giữa token thì OK.',
        color_hex='E0F2FE',
        border_color=ACCENT_BLUE
    )

    add_heading(doc, 'Lỗi 2: Gõ thiếu/thừa dấu ngoặc', level=2, color=ACCENT_RED)
    add_para(doc, 'SAI:', bold=True)
    add_bullet(doc, '{contract_no} (chỉ 1 ngoặc)')
    add_bullet(doc, '{{contract_no} (thiếu 1 ngoặc cuối)')
    add_bullet(doc, '{{ contract_no }} (có dấu cách thừa)')
    add_para(doc, 'ĐÚNG: {{contract_no}} — chính xác 2 mở + 2 đóng + không cách', bold=True, color=ACCENT_GREEN)

    add_heading(doc, 'Lỗi 3: Sai tên biến', level=2, color=ACCENT_RED)
    add_para(doc, 'SAI:', bold=True)
    add_bullet(doc, '{{contract-no}} (gạch ngang thay vì gạch dưới)')
    add_bullet(doc, '{{contractno}} (thiếu gạch dưới)')
    add_bullet(doc, '{{Contract_No}} (viết hoa)')
    add_para(doc, 'ĐÚNG: Chính xác như bảng ở Bước 3. Tất cả chữ thường + gạch dưới _.', bold=True, color=ACCENT_GREEN)

    add_heading(doc, 'Lỗi 4: Copy-paste từ chỗ khác mang theo format ẩn', level=2, color=ACCENT_RED)
    add_para(doc, 'Khi copy {{contract_no}} từ file/website khác paste vào Word, có thể mang theo định dạng vô hình làm Word chia token thành nhiều "run" XML → ERP không nhận.')
    add_callout(doc,
        '✅ Cách an toàn',
        'Gõ TAY 6 token trực tiếp trong Word, không copy-paste. Nếu phải copy-paste: Paste Special → Unformatted Text (Ctrl+Alt+V → chọn "Unformatted Text").',
        color_hex='DCFCE7',
        border_color=ACCENT_GREEN
    )

    add_heading(doc, 'Lỗi 5: Để sót highlight vàng', level=2, color=ACCENT_RED)
    add_para(doc, 'Sau khi gõ token, kiểm tra lại file — không được còn vùng vàng nào nữa. Nếu còn → ERP có thể fill nhầm chỗ.')

    # ─── 6. Test ───
    add_heading(doc, 'Test sau khi sửa', level=1)
    add_para(doc, 'Sau khi sửa xong 4 template, chị test thử 1 HĐ:')
    add_numbered(doc, 'Mở HA_template_SC_FOB.docx đã sửa')
    add_numbered(doc, 'Save as → đổi tên thành TEST_SC_FOB.docx (giữ template gốc nguyên)')
    add_numbered(doc, 'Custom thông tin buyer/grade/giá như HĐ thật')
    add_numbered(doc, 'Upload lên ERP test (HĐ thử) ở /sales/orders/new')
    add_numbered(doc, 'Gọi Phú LV bấm Auto-fill → check xem 6 vùng có fill đúng không')
    add_para(doc, 'Nếu OK → dùng cho production. Nếu sai chỗ nào → báo Minh để check token nào lệch.')

    # ─── 7. Bảng tổng kết ───
    add_heading(doc, 'Bảng tổng kết 6 token (in ra dán cạnh bàn)', level=1)
    add_code_block(doc, '''┌─────────────────────────────────────────────────────┐
│ TOKEN HĐ — Docs điền vào file template Word         │
├─────────────────────────────────────────────────────┤
│ Số HĐ:               {{contract_no}}                │
│ Tên TK ngân hàng:    {{bank_account_name}}          │
│ Số tài khoản:        {{bank_account_no}}            │
│ Tên đầy đủ NH:       {{bank_full_name}}             │
│ Địa chỉ ngân hàng:   {{bank_address}}               │
│ SWIFT code:          {{bank_swift}}                 │
├─────────────────────────────────────────────────────┤
│ Tất cả chữ thường, gạch dưới _,                     │
│ chính xác 2 ngoặc {{ mở + }} đóng,                  │
│ không có dấu cách thừa.                             │
└─────────────────────────────────────────────────────┘''')

    # ─── 8. FAQ ───
    add_heading(doc, 'FAQ', level=1)

    add_heading(doc, 'Q: Có thay đổi nội dung HĐ không?', level=3)
    add_para(doc, 'A: Không. Chỉ đổi 6 vùng để trống thành token. Tất cả nội dung khác (điều khoản, giá, format, font) giữ nguyên y nguyên.')

    add_heading(doc, 'Q: Sau này thêm bank mới (vd. Eximbank) thì có phải sửa template không?', level=3)
    add_para(doc, 'A: Không. Bank info chứa trong token {{bank_*}}, ERP sẽ fill bất cứ bank nào Phú LV chọn. Template không cần biết bank nào.')

    add_heading(doc, 'Q: HĐ có nhiều phụ lục (Packing List, Annex…) thì sao?', level=3)
    add_para(doc, 'A: Nếu phụ lục có số HĐ → cũng gõ {{contract_no}} vào. Nếu phụ lục không có số HĐ → giữ nguyên file. ERP fill được multi-file (max 10 file/HĐ).')

    add_heading(doc, 'Q: HĐ chưa quyết bank (KH confirm sau) thì sao?', level=3)
    add_para(doc, 'A: Phú LV vẫn upload + duyệt được. Token bank sẽ render thành "(See Commercial Invoice)" cho SC hoặc để trống cho PI. Khi tài vụ chốt bank → Phú LV mở HĐ → fill bank → duyệt revision mới.')

    add_heading(doc, 'Q: Em đã quen với highlight vàng, đổi sang token có khó không?', level=3)
    add_para(doc, 'A: Không khó hơn. Cả 2 cách đều là "đặt placeholder ở vị trí cần điền". Highlight vàng = visual marker, token = text marker. Token có ưu thế là ERP đọc được, highlight thì không.')

    # ─── 9. Liên hệ ───
    add_heading(doc, 'Liên hệ khi có vấn đề', level=1)
    add_bullet(doc, 'Hỏi Minh LD nếu không biết token đúng/sai')
    add_bullet(doc, 'Hỏi Phú LV nếu Auto-fill không hoạt động sau khi sửa template')
    add_bullet(doc, 'Báo Minh LD nếu thấy ý tưởng cải tiến quy trình')

    add_callout(doc,
        '📌 Lưu ý cuối',
        'Sửa 4 file template là việc làm 1 LẦN DUY NHẤT. Sau đó mỗi HĐ mới, '
        'Docs làm y nguyên quy trình hiện tại (mở template → custom → upload). '
        'Khác biệt duy nhất: vùng vàng → token. ERP lo phần còn lại.',
        color_hex='FEF3C7',
        border_color=ACCENT_ORANGE
    )

    return doc


def main():
    doc = build_doc()
    path = ROOT / 'HUONG_DAN_DOCS_TEMPLATE_HOP_DONG.docx'
    doc.save(str(path))
    print(f'OK Generated: {path.name} ({path.stat().st_size // 1024} KB)')


if __name__ == '__main__':
    main()
