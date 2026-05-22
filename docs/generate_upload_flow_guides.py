"""
Generate 2 file DOCX huong dan upload flow:
  1. HUONG_DAN_DOCS_TAO_HD_BAN.docx (Nhung + Phuong Anh)
  2. HUONG_DAN_PHU_DUYET_HD.docx (Phu + Minh + Lieu + Minh Anh)

Chay: python docs/generate_upload_flow_guides.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
PRIMARY = RGBColor(0x1B, 0x4D, 0x3E)
ACCENT_RED = RGBColor(0xDC, 0x26, 0x26)
ACCENT_ORANGE = RGBColor(0xD9, 0x77, 0x06)
ACCENT_BLUE = RGBColor(0x1D, 0x4E, 0xD8)
ACCENT_GREEN = RGBColor(0x16, 0xA3, 0x4A)
GRAY = RGBColor(0x6B, 0x72, 0x80)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    if color:
        run.font.color.rgb = color
    elif level == 1:
        run.font.color.rgb = PRIMARY
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_callout(doc, title, body, color_hex='FEF3C7', border_color=ACCENT_ORANGE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_hex)
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    run_t = p_title.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = border_color
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_before = Pt(2)
    run_b = p_body.add_run(body)
    run_b.font.size = Pt(10)
    doc.add_paragraph()


def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


def add_numbered(doc, text):
    return doc.add_paragraph(text, style='List Number')


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1B4D3E')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, 'F9FAFB')
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F3F4F6')
    p = cell.paragraphs[0]
    for i, line in enumerate(code.split('\n')):
        if i > 0:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_cover(doc, title, subtitle, role, date):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(72)
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(8)
    run = p_sub.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(36)
    run = p_meta.add_run(f'Đối tượng: {role}')
    run.bold = True
    run.font.size = Pt(12)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(4)
    run = p_date.add_run(f'Cập nhật: {date}')
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(100)
    run = p_foot.add_run('HUY ANH RUBBER COMPANY LIMITED')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = PRIMARY

    p_foot2 = doc.add_paragraph()
    p_foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_foot2.add_run('ERP System · huyanhrubber.vn')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# GUIDE 1: DOCS — TẠO HĐ BÁN
# ═══════════════════════════════════════════════════════════════════════════════

def build_docs_guide():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    add_cover(doc,
        title='HƯỚNG DẪN TẠO HỢP ĐỒNG BÁN',
        subtitle='Upload flow với template Auto-fill (Phase B)',
        role='DOCS — Nhung + Phương Anh',
        date='2026-05-22'
    )

    # ─── Tổng quan ───
    add_heading(doc, 'Tổng quan workflow', level=1)
    add_para(doc, 'Chị chỉ làm bước 1-3. Khi xong, hệ thống tự chuyển HĐ qua Phú.')

    add_code_block(doc, '''Docs (chị)         →   Phú LV         →   Trung/Huy
─────────────────       ───────────────     ─────────────
1. Tạo đơn ERP          4. Auto-fill       6. In + ký + đóng dấu
2. Upload template      5. Duyệt           7. Gửi KH ký
3. Submit                                  8. Nhận FINAL → upload''')

    # ─── Bước 1: Vào trang ───
    add_heading(doc, 'Bước 1: Vào trang Tạo đơn', level=1)
    add_para(doc, 'URL: huyanhrubber.vn/sales/orders/new', bold=True)
    add_para(doc, 'Hoặc: Menu sidebar → Đơn hàng bán → Đơn hàng → + Tạo mới')

    # ─── Bước 2: Thông tin HĐ ───
    add_heading(doc, 'Bước 2: Điền thông tin Hợp đồng', level=1)
    add_heading(doc, 'Card "Thông tin Hợp đồng"', level=2)

    add_table(doc,
        headers=['Trường', 'Điền gì', 'Bắt buộc?'],
        rows=[
            ['Khách hàng (Buyer)', 'Chọn KH từ dropdown. Chưa có → "+ Thêm KH mới" (điền address + phone)', '✅'],
            ['Số HĐ', 'BỎ TRỐNG — Phú sẽ điền', '❌'],
            ['Ngày HĐ', 'Mặc định hôm nay. Sửa nếu HĐ ký ngày khác', '❌'],
            ['PO# khách hàng', 'Số PO bên KH nếu có (vd APL-2026-001)', '❌'],
        ],
        col_widths=[4, 10, 2]
    )

    add_callout(doc,
        '💡 Smart prefill',
        'Khi chọn KH có đơn cũ → hệ thống tự fill Incoterm + POL + POD + Grade + '
        'Đóng gói + Payment từ đơn gần nhất → chỉ cần verify/sửa nếu khác.',
        color_hex='E0F2FE', border_color=ACCENT_BLUE
    )

    # ─── Bước 3: Sản phẩm & Giá ───
    add_heading(doc, 'Bước 3: Điền Sản phẩm & Giá', level=1)
    add_para(doc, 'Đa số HĐ chỉ có 1 sản phẩm. HĐ nhiều mặt hàng → bấm "+ Thêm SP".')

    add_table(doc,
        headers=['Trường', 'Điền gì', 'Required'],
        rows=[
            ['* Grade', 'SVR3L, SVR10, SVR_CV50, RSS3, SBR1502... (dropdown gợi ý hoặc tự gõ)', '✅'],
            ['* Tấn', 'Tổng tấn HĐ (VD 42)', '✅'],
            ['* $/tấn', 'Đơn giá USD/MT (VD 2,350)', '✅'],
            ['KG/bành', 'Tích 33.33 HOẶC 35 (hoặc cả 2 nếu multi-bale-weight)', 'default 35'],
            ['Bành/cont', '576 cho 20ft / 1152 cho 40ft (auto-tính)', 'default 576'],
            ['Đóng gói', 'Loose Bale / Wooden Pallet / Plastic Pallet / SW Pallet / Metal Box', 'default Loose'],
            ['Ghi chú bao bì', '"Pallet gỗ fumigation", "Bao PE lót đáy cont", "In logo khách"...', '❌'],
            ['Phương thức TT', 'L/C at sight / L/C UPAS 90 days / CAD 5 days / T/T 30/70 / ...', '❌'],
        ],
        col_widths=[3.5, 10, 2.5]
    )

    add_heading(doc, 'Sidebar tự tính', level=3)
    add_bullet(doc, 'Tổng bành = Tấn × 1000 ÷ KG/bành')
    add_bullet(doc, 'Số cont = Tổng bành ÷ Bành/cont (làm tròn lên)')
    add_bullet(doc, 'Giá trị USD = Tấn × $/tấn')

    # ─── Bước 4: Logistics ───
    add_heading(doc, 'Bước 4: Logistics', level=1)

    add_table(doc,
        headers=['Trường', 'Điền gì'],
        rows=[
            ['Incoterm', 'FOB (KH lo cước) / CIF (HA lo cước + bảo hiểm) / CFR / EXW / DDP'],
            ['Cảng xếp (POL)', 'Đà Nẵng / HCM Cát Lái / HCM Cái Mép / Hải Phòng / Quy Nhơn'],
            ['Cảng đích (POD)', 'Chỉ hiện khi Incoterm ≠ FOB/EXW. VD Shanghai, NHAVA SHEVA, BELAWAN'],
            ['Ngày giao dự kiến', 'Tuỳ chọn — render thành "Time of shipment" trong HĐ. Trống → "TBD"'],
        ],
        col_widths=[5, 11]
    )

    # ─── Bước 5: Optional sections ───
    add_heading(doc, 'Bước 5: (Mở rộng nếu cần)', level=1)
    add_para(doc, 'Các section sau collapsed mặc định:')
    add_bullet(doc, 'Chỉ tiêu kỹ thuật — DRC, Moisture, Dirt, Ash... (KH có yêu cầu riêng → expand)')
    add_bullet(doc, 'Hoa hồng môi giới — % hoặc USD/MT (chỉ deal có môi giới)')
    add_bullet(doc, 'Ghi chú đơn hàng — note nội bộ')

    # ─── Bước 6: Upload ───
    add_heading(doc, 'Bước 6: Upload file template HĐ', level=1)

    add_heading(doc, 'File template — đang dùng tạm', level=2)
    add_para(doc, 'Hiện tại chị dùng 4 file mẫu của ERP ở folder:')
    add_code_block(doc, '''docs/contract-template-samples/
├── sample_SC_FOB.docx   (Sales Contract FOB)
├── sample_SC_CIF.docx   (Sales Contract CIF)
├── sample_PI_FOB.docx   (Proforma Invoice FOB)
└── sample_PI_CIF.docx   (Proforma Invoice CIF)''')

    add_heading(doc, 'Quy tắc chọn file', level=2)
    add_bullet(doc, 'HĐ FOB → upload sample_SC_FOB.docx + sample_PI_FOB.docx (2 file)')
    add_bullet(doc, 'HĐ CIF → upload sample_SC_CIF.docx + sample_PI_CIF.docx (2 file)')
    add_para(doc, 'Tối đa 10 file/HĐ, mỗi file max 20MB. Có thể upload thêm phụ lục/packing list.')

    add_callout(doc,
        '💡 Sau này có thay đổi template',
        'Khi chị có template mới: copy file mẫu trên → sửa nội dung theo template Huy Anh → '
        'GIỮ NGUYÊN các placeholder {{...}}. ERP tự fill 18+ trường vào đó (số HĐ, KH, '
        'grade, qty, price, bank, port, payment, packing,...). Xem file '
        'HUONG_DAN_DOCS_TEMPLATE_HOP_DONG.docx để biết chi tiết.',
        color_hex='FEF3C7', border_color=ACCENT_ORANGE
    )

    add_heading(doc, '19 token ERP sẽ tự fill', level=2)

    add_table(doc,
        headers=['Token', 'Nghĩa', 'Nguồn'],
        rows=[
            ['{{contract_no}}', 'Số HĐ', 'Phú điền lúc duyệt'],
            ['{{contract_date}}', 'Ngày HĐ', 'Auto từ form'],
            ['{{buyer_name}}', 'Tên KH', 'Auto từ customer'],
            ['{{buyer_address}}', 'Địa chỉ KH', 'Auto từ customer'],
            ['{{buyer_phone}}', 'Phone KH', 'Auto từ customer'],
            ['{{grade}}', 'Cấp mủ', 'Auto từ order'],
            ['{{quantity}}', 'Tấn', 'Auto'],
            ['{{unit_price}}', 'Đơn giá', 'Auto'],
            ['{{amount}}', 'Tổng tiền', 'Auto'],
            ['{{amount_words}}', 'Số tiền bằng chữ', 'Auto compute'],
            ['{{incoterm}}', 'FOB/CIF/...', 'Auto'],
            ['{{pol}} + {{pod}}', 'Cảng xếp + đích', 'Auto'],
            ['{{packing_desc}}', 'Mô tả đóng gói', 'Auto'],
            ['{{bales_total}} + {{containers}} + {{cont_type}}', 'Số bành + cont + 20\'/40\'', 'Auto'],
            ['{{shipment_time}}', 'Thời gian giao', 'Auto'],
            ['{{payment}}', 'Phương thức TT', 'Auto'],
            ['{{bank_*}} (5 trường)', 'Bank info', 'Phú chọn'],
        ],
        col_widths=[5, 6, 4]
    )

    # ─── Bước 7: Submit ───
    add_heading(doc, 'Bước 7: Submit — Trình Kiểm tra', level=1)
    add_para(doc, 'Bấm nút lớn xanh đậm: "Upload N file + Trình Kiểm tra (Phú LV)"')

    add_heading(doc, 'Sau khi bấm', level=2)
    add_bullet(doc, 'Toast "Đã trình HĐ cho Phú LV (Kiểm tra) duyệt"')
    add_bullet(doc, 'Redirect về /sales/orders — HĐ mới xuất hiện status "draft"')
    add_bullet(doc, 'Phú LV / Minh / Liễu / Minh Anh nhận email + bell — họ vào duyệt')
    add_bullet(doc, 'Chị có thể đóng tab, đợi Phú duyệt')

    # ─── FAQ ───
    add_heading(doc, 'FAQ', level=1)

    faqs = [
        ('Q: Tôi không biết Số HĐ — có sao không?',
         'A: Không sao. Phú LV sẽ điền lúc duyệt. Cứ để trống.'),
        ('Q: Tôi upload nhầm file/sai template — sửa được không?',
         'A: Được. Khi HĐ còn status "reviewing", Phú có thể trả lại chị → chị sửa + trình lại (revision mới).'),
        ('Q: HĐ có phụ lục / packing list — upload chung không?',
         'A: Có. Upload tối đa 10 file. ERP Auto-fill các {{token}} trong TẤT CẢ files. File không có token (vd packing list) → ERP để nguyên.'),
        ('Q: KH chưa có địa chỉ — sao?',
         'A: Trước khi tạo HĐ, vào /sales/customers → mở KH → điền Address + Phone → save. Nếu không, ô buyer_address trong HĐ sẽ trống.'),
        ('Q: Tôi muốn xem trước HĐ render — được không?',
         'A: Phase này chưa có preview. Sau khi Phú Auto-fill, chị xem được trong tab Hợp đồng của đơn (folder "HĐ gửi KH").'),
        ('Q: Tạo nhầm HĐ → xóa được không?',
         'A: Khi HĐ còn draft thì xóa được qua /sales/orders. Hoặc báo Minh LD xóa giúp.'),
    ]
    for q, a in faqs:
        add_heading(doc, q, level=3)
        add_para(doc, a)

    # ─── Tổng kết ───
    add_heading(doc, 'Tổng kết — workflow Docs', level=1)
    add_code_block(doc, '''1. /sales/orders/new
2. Chọn KH → smart prefill
3. Điền 4 trường required: KH, Grade, Tấn, $/tấn
4. Verify defaults: Ngày HĐ, Incoterm, POL, POD, Payment
5. Upload 2 file template .docx (SC + PI tương ứng FOB/CIF)
6. Submit → đợi Phú duyệt''')
    add_para(doc, 'Thời gian/HĐ: ~1-2 phút sau khi quen.', bold=True, color=ACCENT_GREEN)

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# GUIDE 2: PHÚ — DUYỆT HĐ + AUTO-FILL
# ═══════════════════════════════════════════════════════════════════════════════

def build_phu_guide():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    add_cover(doc,
        title='HƯỚNG DẪN DUYỆT HỢP ĐỒNG',
        subtitle='Auto-fill với template Phase B',
        role='Phú LV / Minh LD / Hồ Thị Liễu / Trần Thị Minh Anh',
        date='2026-05-22'
    )

    # ─── Tổng quan ───
    add_heading(doc, 'Tổng quan workflow', level=1)
    add_para(doc, 'Anh/chị chỉ làm bước 3-6. Khoảng 30 giây/HĐ.')
    add_code_block(doc, '''Docs           →   Anh/chị (Kiểm tra)    →   Trung/Huy
─────────         ─────────────────────       ─────────
1. Tạo đơn        3. Nhận thông báo            7. In + ký + đóng dấu
2. Upload .docx   4. Mở HĐ trong Drawer       8. Gửi KH ký
                  5. Auto-fill                 9. Nhận FINAL upload
                  6. Duyệt + Trình ký''')

    # ─── Bước 1: Thông báo ───
    add_heading(doc, 'Bước 1: Nhận thông báo', level=1)
    add_para(doc, 'Khi Docs trình HĐ mới, anh/chị nhận:')
    add_numbered(doc, 'Bell ERP — chuông góc phải trên có số đỏ')
    add_numbered(doc, 'Email đến hộp thư cá nhân (Subject: "📤 HĐ mới ... cần kiểm tra + nhập bank")')
    add_numbered(doc, '(Tương lai) System message trong tab "Trao đổi" của đơn')
    add_para(doc, 'Click vào bell hoặc email → mở thẳng queue Kiểm tra', bold=True)

    # ─── Bước 2: Queue ───
    add_heading(doc, 'Bước 2: Vào Queue Kiểm tra HĐ', level=1)
    add_para(doc, 'URL: huyanhrubber.vn/sales/contracts/review', bold=True)
    add_para(doc, 'Hoặc: Menu sidebar → Đơn hàng bán → Kiểm tra HĐ (chỉ 4 người trên thấy menu)')

    add_heading(doc, 'Queue list — các cột', level=2)
    add_table(doc,
        headers=['Cột', 'Ý nghĩa'],
        rows=[
            ['Số HĐ', 'Hiện "(chưa có số)" nếu Docs chưa điền + Tag 📎 Upload'],
            ['Khách hàng', 'Tên KH'],
            ['Grade', 'Cấp mủ (RSS_3, SVR_10...)'],
            ['Tấn/Cont', 'Số lượng'],
            ['Giá trị USD', 'Tổng đơn'],
            ['Incoterm', 'FOB/CIF/...'],
            ['Người trình', 'Tên Docs + giờ trình'],
        ],
        col_widths=[3, 13]
    )
    add_para(doc, 'Bấm "Mở" ở row → Drawer review chi tiết bên phải.')

    # ─── Bước 3: Drawer ───
    add_heading(doc, 'Bước 3: Mở Drawer Review', level=1)
    add_para(doc, 'Layout Drawer (xếp từ trên xuống):')
    add_bullet(doc, 'Header: tên HĐ + tag Upload + nút Trả lại / Duyệt + Trình ký')
    add_bullet(doc, 'Alert hướng dẫn 4 bước')
    add_bullet(doc, 'Card 🪄 Auto-fill (vàng cam) — Số HĐ + Bank dropdown')
    add_bullet(doc, '(Collapsed) Xem file Docs upload gốc — chỉ tham khảo')
    add_bullet(doc, 'Card ② File anh đã fill — sau Auto-fill, có "Tải về verify"')
    add_bullet(doc, 'Card Tóm tắt đơn (read-only) — đối chiếu nhanh')

    # ─── Bước 4: Auto-fill ───
    add_heading(doc, 'Bước 4: 🪄 Auto-fill — gõ Số HĐ + chọn Bank', level=1)

    add_heading(doc, '4.1. Số HĐ (BẮT BUỘC)', level=2)
    add_bullet(doc, 'Gõ số HĐ theo format Huy Anh quy định')
    add_bullet(doc, 'VD: HA20260100, HA20260101, LTC2026/HA-BRS05 (format đặc biệt)')
    add_bullet(doc, 'Drawer title live update khi gõ — anh/chị thấy ngay "Review HĐ HA20260100"')

    add_heading(doc, '4.2. Bank nhận tiền (TUỲ CHỌN)', level=2)
    add_para(doc, 'Dropdown với 7 ngân hàng có sẵn:')
    add_bullet(doc, 'Vietin Bank — Hue Branch (default — TK chính)')
    add_bullet(doc, 'Vietcombank — Hue Branch')
    add_bullet(doc, 'BIDV — Hue Branch')
    add_bullet(doc, 'Agribank — Hue Branch')
    add_bullet(doc, 'TP Bank — Hue Branch')
    add_bullet(doc, 'Eximbank — Hue Branch')
    add_bullet(doc, 'UOB — Ho Chi Minh City')

    add_callout(doc,
        '💡 Bank chưa quyết được?',
        'Để trống dropdown → 5 field bank trong file render rỗng. Có thể duyệt revision sau '
        'khi tài vụ chốt TK (Trung/Huy trả lại Phú → Phú chọn bank + Auto-fill lại + duyệt '
        'revision mới). HOẶC SC ghi "Payment as per Commercial Invoice" → bank chỉ ở PI.',
        color_hex='FEF3C7', border_color=ACCENT_ORANGE
    )

    add_heading(doc, '4.3. Bấm "🪄 Auto-fill N file Docs upload"', level=2)
    add_para(doc, 'Đợi ~3-5 giây:')
    add_numbered(doc, 'ERP download từng file Docs upload từ Storage')
    add_numbered(doc, 'Render với 18+ token (số HĐ + 5 bank + 12 auto từ DB: buyer/grade/qty/price/incoterm/POL/POD/...)')
    add_numbered(doc, 'Upload file đã render vào ô ②')
    add_para(doc, '→ Toast: "Đã auto-fill N file. Verify rồi bấm Duyệt + Trình ký."',
             bold=True, color=ACCENT_GREEN)

    # ─── Bước 5: Verify ───
    add_heading(doc, 'Bước 5: Verify file đã fill (RECOMMEND)', level=1)
    add_para(doc, 'Card ② "File anh đã fill" hiện mỗi file với nút "Tải về verify" XANH ĐẬM.')

    add_heading(doc, 'Cách verify', level=2)
    add_numbered(doc, 'Bấm "Tải về verify" → mở file trong Word')
    add_numbered(doc, 'Kiểm tra 18+ vùng đã thay đúng (số HĐ, Buyer, Address, Grade, Qty, Price, Bank, SWIFT, Time of shipment, Beneficiary\'s Bank...)')
    add_numbered(doc, 'Nếu file đúng → quay lại ERP bấm "Duyệt + Trình ký"')
    add_numbered(doc, 'Nếu file SAI → Auto-fill lại (đổi bank/số HĐ) HOẶC Trả lại Sale nếu data đơn sai')

    add_heading(doc, 'Tóm tắt đơn (read-only) — đối chiếu nhanh', level=2)
    add_para(doc, 'Card Tóm tắt phía dưới hiện:')
    add_bullet(doc, 'Số HĐ + Revision')
    add_bullet(doc, 'Khách hàng (name + code)')
    add_bullet(doc, 'Grade + Incoterm (Tag)')
    add_bullet(doc, 'Tấn/Cont + $/MT')
    add_bullet(doc, 'Tổng USD (highlight xanh)')
    add_para(doc, '→ So với file vừa download — phải khớp 100%. Nếu khác → có lỗi data trong sales_order.',
             italic=True)

    # ─── Bước 6: Duyệt ───
    add_heading(doc, 'Bước 6: Duyệt + Trình ký', level=1)
    add_para(doc, 'Bấm nút lớn xanh đậm góc phải trên: "✅ Duyệt + Trình ký"')

    add_heading(doc, 'Modal confirm', level=2)
    add_code_block(doc, '''Duyệt + Trình ký
HĐ HA20260100 sẽ chuyển sang [approved], trình Trung/Huy ký.
Sẽ copy 2 file đã fill (ưu tiên) + file gốc còn lại vào folder "HĐ gửi KH".

[Quay lại sửa]  [Duyệt + Trình ký]''')

    add_heading(doc, 'Sau khi bấm', level=2)
    add_bullet(doc, 'Toast "Đã duyệt HĐ HA20260100 → trình ký"')
    add_bullet(doc, 'Drawer close, HĐ biến mất khỏi queue Kiểm tra')
    add_bullet(doc, 'HĐ status: reviewing → approved')
    add_bullet(doc, 'HĐ vào queue Ký HĐ (/sales/contracts/sign)')
    add_bullet(doc, 'Trung + Huy + Minh + Minh Anh nhận email + bell')
    add_bullet(doc, 'Sale (người tạo) nhận email "HĐ đã được duyệt"')
    add_bullet(doc, 'File đã fill auto-copy vào tab Hợp đồng → folder "HĐ gửi KH"')

    # ─── Bước 7: Trả lại ───
    add_heading(doc, 'Bước 7: (Khi cần) Trả lại Sale', level=1)
    add_heading(doc, 'Khi nào Trả lại?', level=2)
    reasons = [
        'Giá đơn / đơn giá sai',
        'Tên KH / địa chỉ thiếu',
        'Incoterm / POL / POD sai',
        'Đóng gói / Số lượng / bales sai',
        'Điều khoản thanh toán không phù hợp',
        'Thời gian giao hàng không đúng',
        'Điều khoản kèm theo sai',
        'Số HĐ / Ngày HĐ sai',
        'Khác (ghi chi tiết)',
    ]
    for r in reasons:
        add_bullet(doc, r)

    add_heading(doc, 'Cách trả lại', level=2)
    add_numbered(doc, 'Bấm nút "Trả lại" ở header (cạnh "Duyệt + Trình ký")')
    add_numbered(doc, 'Modal mở: chọn lý do (multi-select, 9 nhóm) + ghi chi tiết bắt buộc')
    add_numbered(doc, 'Bấm "Trả lại Sale"')

    add_heading(doc, 'Sau khi Trả lại', level=2)
    add_bullet(doc, 'HĐ status: reviewing → rejected')
    add_bullet(doc, 'Sale (Docs) nhận email + bell với full lý do')
    add_bullet(doc, 'HĐ biến mất khỏi queue Kiểm tra')
    add_bullet(doc, 'Sale vào HĐ → "Sửa & Trình lại" → upload revision mới')

    # ─── Edge cases ───
    add_heading(doc, 'Trường hợp đặc biệt', level=1)

    add_heading(doc, 'TH1: File template chưa có placeholder {{...}}', level=2)
    add_para(doc, 'Triệu chứng: Bấm Auto-fill → toast warning "N file lỗi" + Modal liệt kê file fail.', italic=True)
    add_para(doc, 'Nguyên nhân: Docs upload file dùng template cũ (highlight vàng) thay vì placeholder.', italic=True)
    add_para(doc, 'Cách xử lý fallback:', bold=True)
    add_numbered(doc, 'Mở section "Xem N file Docs upload (gốc)" → bấm "Tải" download file gốc')
    add_numbered(doc, 'Mở Word → fill thủ công 2 vùng highlight (số HĐ + bank)')
    add_numbered(doc, 'Save file')
    add_numbered(doc, 'Vào dropzone ô ② → kéo thả file đã fill (REPLACE bộ cũ)')
    add_numbered(doc, 'Bấm Duyệt + Trình ký')
    add_para(doc, 'Báo Docs sửa template theo file HUONG_DAN_DOCS_TEMPLATE_HOP_DONG.docx.', bold=True)

    add_heading(doc, 'TH2: KH chưa có Address trong DB', level=2)
    add_para(doc, 'Triệu chứng: File HĐ render "ADDRESS:" rỗng sau Auto-fill.', italic=True)
    add_para(doc, 'Cách xử lý:', bold=True)
    add_bullet(doc, 'Trả lại Sale với lý do "Tên / Địa chỉ KH sai"')
    add_bullet(doc, 'HOẶC: anh/chị tự vào /sales/customers update → Auto-fill lại')

    add_heading(doc, 'TH3: Trung/Huy trả lại — Phú phải review lại', level=2)
    add_para(doc, 'Triệu chứng: HĐ approved quay về reviewing với note "Trung/Huy trả lại".', italic=True)
    add_para(doc, 'Cách xử lý:', bold=True)
    add_numbered(doc, 'Mở HĐ trong queue Kiểm tra (re-appear)')
    add_numbered(doc, 'Đọc review_notes (note Trung/Huy viết)')
    add_numbered(doc, 'Sửa Bank / Số HĐ → Auto-fill lại')
    add_numbered(doc, 'Duyệt + Trình ký lần 2')

    add_heading(doc, 'TH4: Auto-fill nhiều lần (đổi bank/số HĐ)', level=2)
    add_para(doc, 'Cứ Auto-fill thoải mái — mỗi lần bấm sẽ REPLACE bộ file ở ô ②. File cũ trong Storage tự xóa.')

    # ─── Tổng kết ───
    add_heading(doc, 'Tổng kết — workflow Phú duyệt', level=1)
    add_code_block(doc, '''1. Nhận bell/email "HĐ mới cần kiểm tra"
2. /sales/contracts/review → Mở HĐ
3. Card 🪄: gõ Số HĐ + chọn Bank
4. Bấm Auto-fill 2 file
5. Bấm "Tải về verify" → kiểm Word file đúng chưa
6. Đối chiếu Tóm tắt đơn (read-only)
7. Bấm "Duyệt + Trình ký" → HĐ chuyển Trung/Huy''')
    add_para(doc, 'Thời gian/HĐ: ~30 giây - 1 phút sau khi quen.', bold=True, color=ACCENT_GREEN)

    # ─── FAQ ───
    add_heading(doc, 'FAQ', level=1)

    faqs = [
        ('Q: 4 người cùng có quyền duyệt — ai vào trước duyệt?',
         'A: Ai vào queue trước thấy HĐ thì duyệt được. Dùng chung queue.'),
        ('Q: Duyệt nhầm — undo được không?',
         'A: Không undo trực tiếp. Phải sang queue Ký HĐ → mở HĐ → "Trả lại Phú LV" với lý do → HĐ quay về queue Kiểm tra.'),
        ('Q: Phải verify file mỗi lần không?',
         'A: Lần đầu NÊN verify. Sau quen + tin tưởng → có thể skip, bấm Duyệt thẳng.'),
        ('Q: Trung/Huy ký xong tôi có nhận thông báo không?',
         'A: Có. Email "🎉 HĐ ... đã ký + đóng dấu" khi Trung/Huy upload bản FINAL.'),
        ('Q: HĐ revision nhiều — sao biết khác gì?',
         'A: Mỗi revision là 1 contract row riêng. review_notes có lý do cũ. Sale phải trình LẠI từ đầu.'),
    ]
    for q, a in faqs:
        add_heading(doc, q, level=3)
        add_para(doc, a)

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    docs_doc = build_docs_guide()
    docs_path = ROOT / 'HUONG_DAN_DOCS_TAO_HD_BAN.docx'
    docs_doc.save(str(docs_path))
    print(f'OK Generated: {docs_path.name} ({docs_path.stat().st_size // 1024} KB)')

    phu_doc = build_phu_guide()
    phu_path = ROOT / 'HUONG_DAN_PHU_DUYET_HD.docx'
    phu_doc.save(str(phu_path))
    print(f'OK Generated: {phu_path.name} ({phu_path.stat().st_size // 1024} KB)')


if __name__ == '__main__':
    main()
